# import os
# import logging
# from datetime import datetime
# from fastapi import FastAPI, HTTPException,APIRouter
# from fastapi.responses import JSONResponse
# from pydantic import BaseModel
# from docx import Document
# import mammoth
# from db_config import get_db_connection

# router = APIRouter()

# logging.basicConfig(filename='app.log', level=logging.INFO)

# # Helper function to write the docx file
# def write_array_to_docx(array, name, doc_id, heading, chapter):
#     try:
#         doc = Document()

#         # Add Heading and Chapter
#         doc.add_paragraph(heading, style='Heading 1')
#         doc.add_paragraph(chapter, style='Heading 2')

#         # Add content to the document
#         for item in array:
#             doc.add_paragraph(item)

#         # Create directories if they don't exist
#         output_path = os.path.join(os.getcwd(), 'output', doc_id, name)
#         os.makedirs(os.path.dirname(output_path), exist_ok=True)

#         doc.save(output_path)

#         log_message = f"File Created: {name}\nPath: {output_path}\nDate and Time: {datetime.now()}\n{'-'*40}\n"
#         logging.info(log_message)

#     except Exception as e:
#         logging.error(f"Error creating .docx file: {e}")

# # Helper function to extract text from docx file
# def extract_text_from_docx(file_path):
#     try:
#         with open(file_path, "rb") as docx_file:
#             result = mammoth.extract_raw_text(docx_file)
#             return result.value
#     except Exception as e:
#         logging.error(f"Error extracting text from file: {e}")
#         return ""

# # Route to handle API requests

# @router.get("/process_document")
# async def process_document(doc_id: str):
#     try:
#         conn = get_db_connection()
#         if conn is None:
#             raise HTTPException(status_code=500, detail="Database connection error")
        
#         cursor = conn.cursor()
#         cursor.execute("SELECT * FROM row_document WHERE row_doc_id = %s", (doc_id,))
#         row = cursor.fetchone()
#         conn.close()

#         if not row:
#             raise HTTPException(status_code=404, detail="Document not found")

#         file_path = os.path.join(os.getcwd(),'files', row[1])
#         # Extract text from the Word document
#         file_content = extract_text_from_docx(file_path)
#         # print(file_content)
#         chapter = file_content[0] if file_content else 1

#         # Extract Figures and Tables
#         figure_array = [line for line in file_content.split("\n") if line.startswith("Figure")]
#         table_array = [line for line in file_content.split("\n") if line.startswith("Table")]
#         print("table and figures")

#         # Clean up Figures and Tables by removing colons and dots
#         updated_figure_array = [figure.replace(":", "", 1) for figure in figure_array]
#         updated_table_array = [table.replace(":", "", 1) for table in table_array]

#         write_array_to_docx(updated_table_array, "Table.docx", doc_id, "List of Tables", f"Chapter {chapter}")
#         write_array_to_docx(updated_figure_array, "Figure.docx", doc_id, "List of Figures", f"Chapter {chapter}")

#         # Log the creation of the final document
#         folder_url = f"/output/{doc_id}/"
#         conn = get_db_connection()
#         cursor = conn.cursor()
#         cursor.execute('INSERT INTO final_document (row_doc_id, user_id, final_doc_size, final_doc_url, status, creation_date) '
#                        'VALUES (%s, %s, %s, %s, %s, NOW())', (doc_id, row[2], row[3], folder_url, row[4]))  # Adjust indices
#         conn.commit()
#         conn.close()

#         return JSONResponse(content={"success": True, "doc_id": doc_id})

#     except Exception as e:
#         logging.error(f"Error processing document {doc_id}: {e}")
#         raise HTTPException(status_code=500, detail="Internal Server Error")




import os
import logging
from datetime import datetime
from fastapi import FastAPI, HTTPException, APIRouter
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from docx import Document
import mammoth
from db_config import get_db_connection

router = APIRouter()

logging.basicConfig(filename='app.log', level=logging.INFO)

def write_array_to_docx(array, name, doc_id, heading, chapter, username, current_date):
    try:
        doc = Document()

        # Add Heading and Chapter
        doc.add_paragraph(heading, style='Heading 1')
        doc.add_paragraph(chapter, style='Heading 2')

        # Add content to the document
        for item in array:
            doc.add_paragraph(item)

        # Create directories with username, current_date, and doc_id
        output_path = os.path.join(os.getcwd(), 'output', username, current_date, doc_id, name)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        doc.save(output_path)

        log_message = f"File Created: {name}\nPath: {output_path}\nDate and Time: {datetime.now()}\n{'-'*40}\n"
        logging.info(log_message)

    except Exception as e:
        logging.error(f"Error creating .docx file: {e}")

# def extract_text_from_docx(file_path):
#     try:
#         with open(file_path, "rb") as docx_file:
#             result = mammoth.extract_raw_text(docx_file)
#             # print(result.value)
#             return result.value
#     except Exception as e:
#         logging.error(f"Error extracting text from file: {e}")
#         return ""

from docx import Document
# def extract_text_from_docx(file_path):
#     try:
#         doc = Document(file_path)
#         # The paragraphs property only returns paragraphs from the main document body,
#         # which excludes paragraphs that are inside tables.
#         paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
#         print("\n".join(paragraphs))
#         return "\n".join(paragraphs)
#     except Exception as e:
#         logging.error(f"Error extracting text from file: {e}")
#         return ""

import re
def extract_text_tables_figures(file_path):
    try:
        doc = Document(file_path)
        chapters = {}
        current_chapter = None
        table_count = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Detect chapter headings
                if re.match(r'Chapter \d+:', text, re.IGNORECASE):
                    current_chapter = text
                    chapters[current_chapter] = {'paragraphs': [], 'figures': [], 'tables': []}
                    table_count = 0
                elif current_chapter:
                    # Check for figures
                    if text.lower().startswith("figure"):
                        cleaned_text = text.replace(":", "", 1)
                        chapters[current_chapter]['figures'].append(cleaned_text)
                        # chapters[current_chapter]['figures'].append(text)
                    else:
                        chapters[current_chapter]['paragraphs'].append(text)
                    
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Detect chapter headings
                if re.match(r'Chapter \d+:', text, re.IGNORECASE):
                    current_chapter = text
                    # chapters[current_chapter] = {'paragraphs': [], 'figures': [], 'tables': []}
                    table_count = 0
                elif current_chapter:
                    # Check for figures
                    if text.lower().startswith("table"):
                        cleaned_text = text.replace(":", "", 1)
                        chapters[current_chapter]['tables'].append(cleaned_text)
                    else:
                        chapters[current_chapter]['paragraphs'].append(text)

        return chapters

    except Exception as e:
        logging.error(f"Error extracting text from file: {e}")
        return {}



def create_docx(chapters, content_type, doc_id, username, current_date, name):
    doc = Document()
    doc.add_heading(f"List of {content_type}", level=1)

    for chapter, content in chapters.items():
        doc.add_heading(chapter, level=2)
        if content_type.lower() == "figures":
            for figure in content['figures']:
                doc.add_paragraph(figure)
        elif content_type.lower() == "tables":
            for table in content['tables']:
                doc.add_paragraph(table)
                
    output_path = os.path.join(os.getcwd(), 'output', username, current_date, doc_id, 'doc', name)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)



@router.get("/process_document")
async def process_document(doc_id: str):
    try:
        conn = get_db_connection()
        if conn is None:
            raise HTTPException(status_code=500, detail="Database connection error")
        
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM row_document WHERE row_doc_id = %s", (doc_id,))
        row = cursor.fetchone()
        cursor.execute("SELECT admin_name from admins where admin_id = %s",(row[5],))
        user = cursor.fetchone()
        conn.close()
        
        if not row:
            raise HTTPException(status_code=404, detail="Document not found")

        username = user[0]
        # print(username)
        current_date = datetime.now().strftime("%Y-%m-%d")
        file_path = os.path.join(os.getcwd(), 'files', row[1])
        # file_content = extract_text_from_docx(file_path)
        chapters = extract_text_tables_figures(file_path)
        create_docx(chapters, "Figures", doc_id, username, current_date, "Figure.docx")
        create_docx(chapters, "Tables", doc_id, username, current_date, "Table.docx")
        # chapter = file_content[0] if file_content else 1
        
        chapter = 1

        # figure_array = [line for line in file_content.split("\n") if line.startswith("Figure")]
        # table_array = [line for line in file_content.split("\n") if line.startswith("Table")]

        # updated_figure_array = [figure.replace(":", "", 1) for figure in figure_array]
        # updated_table_array = [table.replace(":", "", 1) for table in table_array]

        # # Pass username and current_date to the function
        # write_array_to_docx(updated_table_array, "Table.docx", doc_id, "List of Tables", f"Chapter {chapter}", username, current_date)
        # write_array_to_docx(updated_figure_array, "Figure.docx", doc_id, "List of Figures", f"Chapter {chapter}", username, current_date)

        # Update folder_url to include username and current_date
        folder_url = f"/output/{username}/{current_date}/{doc_id}/"
        
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO final_document 
            (row_doc_id, user_id, final_doc_size, final_doc_url, status, creation_date)
            VALUES (%s, %s, %s, %s, %s, NOW())
        ''', (doc_id, row[2], row[3], folder_url, row[4]))
        conn.commit()
        conn.close()

        return JSONResponse(content={"success": True, "doc_id": doc_id})

    except Exception as e:
        # logging.error(f"Error processing document {doc_id}: {e}")
        print(e)
        raise HTTPException(status_code=500, detail="Internal Server Error")