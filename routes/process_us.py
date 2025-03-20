import re
from docx.shared import RGBColor
from num2words import num2words
from word2number import w2n
import enchant
from fastapi import APIRouter, HTTPException, Depends, Query
from fastapi.responses import FileResponse
import os
import docx
from sqlalchemy import text
from db_config import get_db_connection
import mammoth
from datetime import datetime
from pathlib import Path
import logging  
import roman
from urllib.parse import urlparse
from datetime import datetime, timedelta
from jose import JWTError, jwt
from typing import Dict
from pydantic import BaseModel, RootModel
from process_module.punctuation import process_doc_function1
from process_module.NumberAndScientificUnit import process_doc_function2
from process_module.hyphen import process_doc_function3
from process_module.formatting import process_doc_function4
from process_module.chapters import process_doc_function6
from process_module.heading import process_doc_function7

router = APIRouter()

us_dict = enchant.Dict("en_US")

# us_dict = enchant.DictWithPWL("en_US","mywords.txt")

global_logs = []

def fetch_abbreviation_mappings():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT original_word, abbreviated_form FROM abbreviation_mapping")
    mappings = cursor.fetchall()
    conn.close()
    return {row[0]: row[1] for row in mappings}


century_map = {
    1: "first",
    2: "second",
    3: "third",
    4: "fourth",
    5: "fifth",
    6: "sixth",
    7: "seventh",
    8: "eighth",
    9: "ninth",
    10: "tenth",
    11: "eleventh",
    12: "twelfth",
    13: "thirteenth",
    14: "fourteenth",
    15: "fifteenth",
    16: "sixteenth",
    17: "seventeenth",
    18: "eighteenth",
    19: "nineteenth",
    20: "twentieth",
    21: "twenty-first",
    22: "twenty-second",
    23: "twenty-third",
    24: "twenty-fourth",
    25: "twenty-fifth",
}

def apply_abbreviation_mapping(text, abbreviation_dict, line_number):
    global global_logs
    words = text.split()
    updated_text = []
    for word in words:
        updated_word = abbreviation_dict.get(word, word)
        if word != updated_word:
            global_logs.append(f"[apply_abbreviation_mapping] Line {line_number}: '{word}' -> '{updated_word}'")
        updated_text.append(updated_word)
    return ' '.join(updated_text)



def apply_number_abbreviation_rule(text, line_number):
    """
    Replaces 'Number X' or 'number X' with 'No. X' or 'no. X' and logs changes.
    :param text: The input text.
    :param line_number: Line number for logging.
    :return: Updated text with number abbreviations applied.
    """
    global global_logs

    def replace_number(match):
        word = match.group(1)
        num = match.group(2)
        updated_text = f"No. {num}" if word.istitle() else f"no. {num}"
        global_logs.append(f"[apply_number_abbreviation_rule] Line {line_number}: '{match.group(0)}' -> '{updated_text}'")
        return updated_text

    pattern = r'\b(Number|number)\s(\d+)\b'
    return re.sub(pattern, replace_number, text)


def apply_numerals_rule(text):
    def text_to_num(match):
        try:
            return str(w2n.word_to_num(match.group(0)))
        except ValueError:
            return match.group(0)
    text = re.sub(r'\b(\w+ and a half|\w+ and \w+/\w+)\b', text_to_num, text)
    text = re.sub(r'\b(\w+-\w+/\w+)\b', text_to_num, text)
    text = re.sub(r'\b(\w+) years? old\b', text_to_num, text)
    text = re.sub(r'\b(\w+ (first|second|third|fourth|fifth|sixth|seventh|eighth|ninth|tenth|eleventh|twelfth))\b', text_to_num, text)
    return text

# Done
def replace_percent_with_symbol(text):
    global global_logs

    modified_text = []
    lines = text.splitlines()

    for line_number, line in enumerate(lines, 1):
        # Find matches for numbers followed by 'percent' or 'per cent'
        matches = re.findall(r"(\d+)\s?(percent|per cent)", line, flags=re.IGNORECASE)

        # If there are matches, replace them and store the change in the global log
        if matches:
            for match in matches:
                original_text = f"{match[0]} {match[1]}"
                modified_text_line = line.replace(original_text, f"{match[0]}%")
                global_logs.append(
                    f"[replace_percent_with_symbol] Line {line_number}: {original_text} -> {match[0]}%"
                )
                line = modified_text_line  # Update the line after the change

        modified_text.append(line)  # Add the modified line to the final text

    return "\n".join(modified_text)


# Done
def convert_century(text, line_number_offset):
    """
    Converts century notation like '21st' to 'the twenty-first century'
    and logs the changes with line numbers.
    
    :param text: The entire text to process, possibly spanning multiple lines.
    :param line_number_offset: The starting line number for this chunk of text.
    :return: The updated text with century notations converted.
    """
    global global_logs  # Global log to record changes
    lines = text.split('\n')  # Split text into individual lines
    updated_lines = []

    for index, line in enumerate(lines):
        words = line.split()  # Split line into words
        for i, word in enumerate(words):
            match = re.match(r"(\d+)(st|nd|rd|th)$", word)  # Match century notation
            if match:
                num = int(match.group(1))
                if num in century_map:
                    # Original and converted word
                    original_word = match.group(0)
                    converted_word = f"the {century_map[num]} century"
                    
                    # Log the change with the actual line number
                    global_logs.append(
                        f"[convert century] Line {line_number_offset + index}: {original_word} -> {converted_word}"
                    )
                    
                    # Replace the word in the line
                    words[i] = converted_word
        
        # Rebuild the updated line
        updated_lines.append(' '.join(words))

    # Return the updated text with all lines rebuilt
    return '\n'.join(updated_lines)



def clean_word(word):
    return word.strip(",.?!:;\"'()[]{}")

# def clean_word(word):
#     return word

    
    

def replace_straight_quotes_with_curly(text):
    # Replace straight double quotes with opening and closing curly quotes
    text = re.sub(r'(^|[\s([{])"', r'\1“', text)  # Opening double quotes
    text = re.sub(r'"', r'”', text)
    
    # Replace straight single quotes with opening and closing curly quotes
    text = re.sub(r"(^|[\s([{])'", r'\1‘', text)  # Opening single quotes
    text = re.sub(r"'", r'’', text)  # Closing single quotes
    
    text = re.sub(r"([a-zA-Z]+)'([a-zA-Z]+)", r"\1‘\2", text)  # Curly starting single quote after word
    
    return text



# Done
def correct_acronyms(text, line_number):
    global global_logs
    original_text = text
    words = text.split()
    corrected_words = []
    for word in words:
        original_word = word
        if re.match(r"([a-z]\.){2,}[a-z]\.?", word):
            word = word.replace(".", "")
        elif re.match(r"([A-Z]\.){2,}[A-Z]\.?", word):
            word = word.replace(".", "")
        if word != original_word:
            global_logs.append(
                f"[correct_acronyms] Line {line_number}: '{original_word}' -> '{word}'"
            )
        corrected_words.append(word)
    corrected_text = " ".join(corrected_words)
    return corrected_text



def enforce_am_pm(text, line_num):
    """
    Ensures consistent formatting for 'am' and 'pm' in the entire paragraph and logs changes.
    :param text: The paragraph text to process.
    :param line_num: The line number in the document for logging.
    :return: The updated text with corrected 'am' and 'pm' formats.
    """
    global global_logs  # Use a global log to record changes
    original_text = text  # Store the original text for comparison
    words = text.split()  # Split the paragraph into words

    corrected_words = []
    for word in words:
        original_word = word  # Store the original word for logging
        word_lower = word.lower()  # Convert word to lowercase for comparison

        # Check and correct 'am' or 'pm' formats
        if word_lower in {"am", "a.m", "pm", "p.m"}:
            if "a" in word_lower:
                corrected_word = "a.m."
            elif "p" in word_lower:
                corrected_word = "p.m."
            
            # Log the change if the word was modified
            if corrected_word != original_word:
                global_logs.append(
                    f"[am pm change] Line {line_num}: '{original_word}' -> '{corrected_word}'"
                )
        else:
            corrected_word = word  # Keep the word unchanged if no match

        corrected_words.append(corrected_word)  # Add the corrected word to the list

    # Join the corrected words to form the updated paragraph
    corrected_text = " ".join(corrected_words)

    return corrected_text



# Done
# [apostrophes change] : 60's -> 1960s 
def remove_unnecessary_apostrophes(word, line_num):
    original_word = word
    global global_logs
    word = re.sub(r"(\d{4})'s\b", r"\1s", word)
    word = re.sub(r"'(\d{2})s\b", r"\1s", word)
    word = re.sub(r"(\d{4}s)'\b", r"\1", word)
    word = re.sub(r"(\d+)'(s|st|nd|rd|th)\b", r"\1\2", word)
    word = re.sub(r"^(\d{2})s\b", r"19\1s", word)
    if word != original_word:
        global_logs.append(f"[apostrophes change] Line {line_num}: {original_word} -> {word}")
    
    return word


# Not working
def spell_out_number_and_unit_with_rules(sentence, line_number):
    global global_logs
    original_words = sentence.split()
    modified_words = original_words[:]
    unit_pattern = r"(\d+)\s+([a-zA-Z]+)"
    number_pattern = r"\b(\d+)\b"

    for i, word in enumerate(original_words):
        # Handle number followed by unit
        if re.match(unit_pattern, " ".join(original_words[i:i+2])):
            continue  # Skip since it's already formatted correctly
        # Spell out numbers less than 10
        elif re.match(number_pattern, word):
            number = int(word)
            if number < 10:
                modified_words[i] = num2words(number, to="cardinal")
    
    # Log only changes
    for orig, mod in zip(original_words, modified_words):
        if orig != mod:
            global_logs.append(f"[spell_out_number_and_unit_with_rules] Line {line_number}: '{orig}' -> '{mod}'")
    return " ".join(modified_words)



def use_numerals_with_percent(text):
    global global_logs

    lines = text.splitlines()
    modified_text = []

    for line_number, line in enumerate(lines, 1):
        original_line = line
        modified_line = line
        def replace_spelled_out_percent(match):
            word = match.group(1)
            try:
                num = w2n.word_to_num(word.lower())
                modified = f"{num}%"
                global_logs.append(
                    f"[numerals with percent] Line {line_number}: '{word} percent' -> '{modified}'"
                )
                return modified
            except ValueError:
                return match.group(0)

        modified_line = re.sub(
            r"\b([a-zA-Z\s\-]+)\s?(percent|per cent|percentage)\b",
            replace_spelled_out_percent,
            modified_line,
            flags=re.IGNORECASE,
        )

        def replace_numerical_percent(match):
            number = match.group(1)
            modified = f"{number}%"
            global_logs.append(
                f"[numerals with percent] Line {line_number}: '{match.group(0)}' -> '{modified}'"
            )
            return modified

        modified_line = re.sub(
            r"(\d+)\s?(percent|per cent|percentage)\b", replace_numerical_percent, modified_line, flags=re.IGNORECASE
        )

        modified_text.append(modified_line)

    return "\n".join(modified_text)




def enforce_eg_rule_with_logging(text):
    lines = text.splitlines()
    updated_lines = []
    for line_number, line in enumerate(lines, start=1):
        original_line = line

        # Step 1: Match "eg" or "e.g." with optional surrounding spaces and punctuation
        new_line = re.sub(r'\beg\b', 'e.g.', line, flags=re.IGNORECASE)
        new_line = re.sub(r'\beg,\b', 'e.g.', new_line, flags=re.IGNORECASE)  # Handle "eg,"

        # Step 2: Fix extra periods like `e.g..` or `e.g...,` and ensure proper punctuation
        new_line = re.sub(r'\.([.,])', r'\1', new_line)  # Removes an extra period before a comma or period
        new_line = re.sub(r'\.\.+', '.', new_line)  # Ensures only one period after e.g.

        # Step 3: Remove comma if e.g... is followed by it (e.g..., -> e.g.)
        new_line = re.sub(r'e\.g\.,', 'e.g.', new_line)

        # Step 4: Change e.g, to e.g.
        new_line = re.sub(r'e\.g,', 'e.g.', new_line)

        # Log changes if the line is updated
        if new_line != line:
            global_logs.append(
                f"[e.g. correction] Line {line_number}: {line.strip()} -> {new_line.strip()}"
            )
        
        updated_lines.append(new_line)
    return "\n".join(updated_lines)




def enforce_ie_rule_with_logging(text):
    lines = text.splitlines()
    updated_lines = []
    for line_number, line in enumerate(lines, start=1):
        original_line = line

        # Step 1: Match "ie" or "i.e." with optional surrounding spaces and punctuation
        new_line = re.sub(r'\bie\b', 'i.e.', line, flags=re.IGNORECASE)  # Handle standalone "ie"
        new_line = re.sub(r'\bie,\b', 'i.e.', new_line, flags=re.IGNORECASE)  # Handle "ie,"

        # Step 2: Fix extra periods like `i.e..` or `i.e...,` and ensure proper punctuation
        new_line = re.sub(r'\.([.,])', r'\1', new_line)  # Removes an extra period before a comma or period
        new_line = re.sub(r'\.\.+', '.', new_line)  # Ensures only one period after i.e.

        # Step 3: Remove comma if i.e... is followed by it (i.e..., -> i.e.)
        new_line = re.sub(r'i\.e\.,', 'i.e.', new_line)
        
        # Step 4: Change i.e, to i.e.
        new_line = re.sub(r'i\.e,', 'i.e.', new_line)

        # Log changes if the line is updated
        if new_line != line:
            global_logs.append(
                f"[i.e. correction] Line {line_number}: {line.strip()} -> {new_line.strip()}"
            )
        
        updated_lines.append(new_line)
    return "\n".join(updated_lines)





def standardize_etc(text):
    lines = text.splitlines()
    updated_lines = []
    pattern = r'\b(e\.?tc|e\.t\.c|e\.t\.c\.|et\.?\s?c|et\s?c|etc\.?|etc|et cetera|etcetera|Etc\.?|Etc|‘and etc\.’|et\.?\s?cetera|etc\.?,?|etc\.?\.?|etc\,?\.?)\b'
    
    for line_number, line in enumerate(lines, start=1):
        original_line = line
        
        # Replace all matches of "etc." variations with "etc."
        new_line = re.sub(pattern, 'etc.', line, flags=re.IGNORECASE)
        
        # Explicitly replace "etc.." with "etc."
        new_line = re.sub(r'etc\.\.+', 'etc.', new_line)
        
        # Explicitly replace "etc.." with "etc."
        new_line = re.sub(r'etc\.\.+', 'etc.', new_line)
        
        # Explicitly replace "etc.," with "etc."
        new_line = re.sub(r'etc\.,', 'etc.', new_line)

        # Log changes if the line is updated
        if new_line != line:
            global_logs.append(f"[etc. correction] Line {line_number}: {line.strip()} -> {new_line.strip()}")
        
        updated_lines.append(new_line)
    return "\n".join(updated_lines)



# def adjust_ratios(text):
#     return re.sub(r"(\d)\s*:\s*(\d)", r"\1 : \2", text)



def adjust_ratios(text):
    """
    Ensures proper formatting of ratios with spaces around the colon (e.g., "1:2" -> "1 : 2").

    Args:
        text (str): Input text to process.

    Returns:
        str: Updated text.
    """
    global global_logs
    def process_ratio(match):
        original = match.group(0)
        modified = f"{match.group(1)} : {match.group(2)}"
        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[adjust_ratios] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified
    return re.sub(r"(\d)\s*:\s*(\d)", process_ratio, text)





def correct_chapter_numbering(text, chapter_counter):
    chapter_pattern = re.compile(r"(?i)\bchapter\s+((?:[IVXLCDM]+)|(?:[a-z]+)|(?:\d+))[:.]?\s")
    def replace_chapter_heading(match):
        chapter_content = match.group(1)
        if re.match(r"^[IVXLCDM]+$", chapter_content, re.IGNORECASE):
            chapter_number = roman.fromRoman(chapter_content.upper())
        elif re.match(r"^[a-z]+$", chapter_content, re.IGNORECASE):
            chapter_number = w2n.word_to_num(chapter_content.lower())
        else:
            chapter_number = int(chapter_content)
        return f"Chapter {chapter_number}: "
    return chapter_pattern.sub(replace_chapter_heading, text)


def enforce_number_spelling_rule(text: str):
    num_to_words = {
        "1": "one", "2": "two", "3": "three", "4": "four", "5": "five",
        "6": "six", "7": "seven", "8": "eight", "9": "nine"
    }
    units = r"(kg|g|mg|cm|mm|km|m|l|ml|%)"
    sentences = re.split(r"(?<=[.!?])\s+", text)
    updated_sentences = []
    for sentence in sentences:
        numbers = re.findall(r"\b\d+\b", sentence)
        if any(int(num) >= 10 for num in numbers) and any(int(num) < 10 for num in numbers):
            updated_sentences.append(sentence)
            continue
        def replace_number(match):
            number = match.group()
            if number in num_to_words:
                if re.search(rf"\b{number}\b\s+{units}", sentence):
                    return number
                if re.search(rf"\b{number}-[a-zA-Z-]+", sentence):
                    return num_to_words[number]
                return num_to_words[number]
            return number
        updated_sentence = re.sub(r"\b\d+\b", replace_number, sentence)
        updated_sentences.append(updated_sentence)
    return " ".join(updated_sentences)




# Done
# [insert_thin_space_between_number_and_unit] Line 31: '5kg' -> '5 kg'
def insert_thin_space_between_number_and_unit(text, line_number):
    global global_logs
    original_text = text
    thin_space = '\u2009'
    
    pattern = r"(\d+)(?=\s?[a-zA-Z]+)(?!\s?°)"

    updated_text = text  # Initialize updated text to the original

    matches = re.finditer(pattern, text)
    for match in matches:
        number = match.group(1)  # This is the number
        unit_start = match.end()
        unit = text[unit_start:].split()[0] 
        
        original_word = number + unit
        updated_word = number + thin_space + unit

        updated_text = updated_text.replace(original_word, updated_word, 1)

        global_logs.append(
            f"[insert_thin_space_between_number_and_unit] Line {line_number}: '{original_word}' -> '{updated_word}'"
        )
    return updated_text




# def format_dates(text):
#     text = re.sub(r"\b(\d+)\s?(BCE|CE)\b", lambda m: f"{m.group(1)} {m.group(2).lower()}", text)
#     text = re.sub(r"\b(AD|BC)\.\b", r"\1 ", text)
#     text = re.sub(r"(\d+)\s?(BCE|CE|AD|BC)\b", r"\1 \2", text)
#     return text


# Done
# [format_dates] Line 5: '386 BCE' -> '386 bce'
def format_dates(text, line_number):
    global global_logs

    def log_and_replace(pattern, replacement, text):
        def replacer(match):
            original = match.group(0)
            updated = replacement(match)
            if original != updated:
                global_logs.append(
                    f"[format_dates] Line {line_number}: '{original}' -> '{updated}'"
                )
            return updated
        return re.sub(pattern, replacer, text)
    text = log_and_replace(
        r"\b(\d+)\s?(BCE|CE)\b",
        lambda m: f"{m.group(1)} {m.group(2).lower()}",
        text
    )
    text = log_and_replace(
        r"\b(AD|BC)\.\b",
        lambda m: f"{m.group(1)} ",
        text
    )
    text = log_and_replace(
        r"(\d+)\s?(BCE|CE|AD|BC)\b",
        lambda m: f"{m.group(1)} {m.group(2)}",
        text
    )
    return text


# Done
# [remove_space_between_degree_and_direction] Line 10: '52 °N' -> '52°N'
def remove_space_between_degree_and_direction(text, line_number):
    global global_logs
    pattern = r"(\d+) \s*[º°]\s*(N|S|E|W)\b"
    def log_replacement(match):
        original_text = match.group(0)
        updated_text = match.group(1) + "º" + match.group(2)
        global_logs.append(
            f"[remove_space_between_degree_and_direction] Line {line_number}: '{original_text}' -> '{updated_text}'"
        )
        return updated_text
    updated_text = re.sub(pattern, log_replacement, text)
    return updated_text



# Done
# km not Km; kg not Kg; l not L. (2.9)
def enforce_lowercase_units(text, line_number):
    global global_logs
    unit_patterns = [
        (r"(\d+)\s*(K)(m|g|l)", 'K', 'k'),
        (r"(\d+)\s*(G)(m)", 'G', 'g'),
        (r"(\d+)\s*(M)(g)", 'M', 'm'),
        (r"(\d+)\s*(T)(g)", 'T', 't'),
        (r"(\d+)\s*(L)\b", 'L', 'l'),
        (r"(\d+)\s*(M)\b", 'M', 'm'),
        (r"(\d+)\s*(kg|mg|g|cm|m|km|l|s|h|min)", r"\1 \2", None)
    ]
    updated_text = text
    for pattern, original, updated in unit_patterns:
        matches = re.finditer(pattern, updated_text)
        for match in matches:
            original_text = match.group(0)
            if updated is not None:
                updated_text = updated_text.replace(original_text, original_text.replace(original, updated))
                global_logs.append(
                    f"[enforce_lowercase_units] Line {line_number}: '{original_text}' -> '{original_text.replace(original, updated)}'"
                )
            else:
                updated_text = updated_text.replace(original_text, f"{match.group(1)} {match.group(2)}")
                global_logs.append(
                    f"[enforce_lowercase_units] Line {line_number}: '{original_text}' -> '{match.group(1)} {match.group(2)}'"
                )
    return updated_text


# Done
# [precede_decimal_with_zero] Line 22: '.76' -> '0.76'
def precede_decimal_with_zero(text, line_number):
    global global_logs
    pattern = r"(?<!\d)(?<!\d\.)\.(\d+)"
    def log_replacement(match):
        original_text = match.group(0)
        updated_text = "0." + match.group(1)
        global_logs.append(
            f"[precede_decimal_with_zero] Line {line_number}: '{original_text}' -> '{updated_text}'"
        )
        return updated_text
    updated_text = re.sub(pattern, log_replacement, text)
    return updated_text


# Done
def adjust_terminal_punctuation_in_quotes(text):
    text = re.sub(
        r"([‘“])([^’”]*[?!])([’”])\.",
        r"\1\2\3",
        text
    )
    return text




def enforce_serial_comma(text):
    lines = text.splitlines()
    updated_lines = []

    for line_number, line in enumerate(lines, start=1):
        original_line = line

        # Add a comma before "and" or "or" in lists
        new_line = re.sub(
            r'([^,]+), ([^,]+) (or) ([^,]+)',
            r'\1, \2, \3 \4',
            line
        )
        # Explicitly handle cases where "or" does not get the serial comma
        new_line = re.sub(
            r'([^,]+), ([^,]+) (and) ([^,]+)',
            r'\1, \2, \3 \4',
            new_line
        )
        if new_line != line:
            global_logs.append(f"[Serial comma correction] Line {line_number}: {line.strip()} -> {new_line.strip()}")
        
        updated_lines.append(new_line)
    return "\n".join(updated_lines)



def correct_possessive_names(text, line_number):
    global global_logs
    pattern_singular_possessive = r"\b([A-Za-z]+s)\b(?<!\bs')'"
    matches_singular = re.finditer(pattern_singular_possessive, text)
    updated_text = text
    for match in matches_singular:
        original_text = match.group(0)
        updated_text_singular = match.group(1)[:-1] + "'s"
        updated_text = updated_text.replace(original_text, updated_text_singular)
        global_logs.append(
            f"[correct_possessive_names] Line {line_number}: '{original_text}' -> '{updated_text_singular}'"
        )
    pattern_plural_possessive = r"\b([A-Za-z]+s)'\b"
    matches_plural = re.finditer(pattern_plural_possessive, updated_text)
    for match in matches_plural:
        original_text = match.group(0)
        updated_text_plural = match.group(1) + "'"
        updated_text = updated_text.replace(original_text, updated_text_plural)
        global_logs.append(
            f"[correct_possessive_names] Line {line_number}: '{original_text}' -> '{updated_text_plural}'"
        )
    return updated_text




# Done
# http://www.PHi.com/authorguidelines not http://www.PHi.com/authorguidelines/
def remove_concluding_slashes_from_urls(text, line_number):
    global global_logs
    pattern = r"(https?://[^\s/]+(?:/[^\s/]+)*)/"
    matches = re.finditer(pattern, text)
    updated_text = text
    
    for match in matches:
        original_text = match.group(0)
        updated_text_url = match.group(1)  # URL without the concluding slash (e.g., "https://example.com")
        updated_text = updated_text.replace(original_text, updated_text_url)
        
        # Log the change
        global_logs.append(
            f"[remove_concluding_slashes_from_urls] Line {line_number}: '{original_text}' -> '{updated_text_url}'"
        )    
    return updated_text



# def clean_web_addresses(text):
#     return re.sub(r"<(https?://[^\s<>]+)>", r"\1", text)


def clean_web_addresses(text):
    """
    Removes angle brackets around web addresses (e.g., "<http://example.com>" -> "http://example.com").
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs
    def process_web_address(match):
        original = match.group(0)
        modified = match.group(1)

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[clean_web_addresses] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified
    return re.sub(r"<(https?://[^\s<>]+)>", process_web_address, text)



def format_ellipses_in_series(text):
    # Matches series like "x1, x2, ..., xn" and ensures the ellipsis has a comma and space after it.
    text = re.sub(r"(\w+),\s*(\w+),\s*\.\.\.\s*(\w+)", r"\1, \2, …, \3", text)
    return text



def format_chapter_title(text):
    match = re.match(r"Chapter\s+([\dIVXLCDM]+)[\.:]\s*(.*)", text, re.IGNORECASE)
    if match:
        chapter_number = match.group(1)
        chapter_title = match.group(2).rstrip('.')
        words = chapter_title.split()
        formatted_title = " ".join([
            word.capitalize() if i == 0 or len(word) >= 4 else word.lower()
            for i, word in enumerate(words)
        ])
        # print(formatted_title)
        return f"{chapter_number}. {formatted_title}"
    return text




def format_titles_us_english_with_logging(text, doc_id):
    global global_logs
    titles = {
        "doctor": "Dr.",
        "mister": "Mr.",
        "misses": "Mrs.",
        "miss": "Miss.",
        "ms": "Ms.",
        "professor": "Professor",
        "sir": "Sir",
        "madam": "Madam",
        "saint": "St",
    }    
    lines = text.splitlines()
    updated_lines = []
    for line_number, line in enumerate(lines, start=1):
        original_line = line
        for title, replacement in titles.items():
            new_line = re.sub(rf"\b{title}\b", replacement, line, flags=re.IGNORECASE)
            if new_line != line:
                global_logs.append(f"[shorten title] Line {line_number}: {title} -> {replacement}")
                line = new_line
        updated_lines.append(line)
    return "\n".join(updated_lines)


def units_with_bracket(text, doc_id):
    units = {
        "s": "second",
        "m": "meter",
        "kg": "kilogram",
        "A": "ampere",
        "K": "kelvin",
        "mol": "mole",
        "cd": "candela"
    }
    used_units = set()
    global global_logs
    processed_lines = []
    for line_num, line in enumerate(text.splitlines(), start=1):
        def replace_unit(match):
            number = match.group(1)
            unit = match.group(2)
            if unit in used_units:
                return match.group(0)
            else:
                used_units.add(unit)
                full_form = units[unit]
                if unit != "mol" and not full_form.endswith("s"):
                    full_form += "s"
                replacement = f"{number} {full_form} ({unit.lower()})"
                global_logs.append(
                    f"Line {line_num}: {match.group(0)} -> {replacement}"
                )
                return replacement
        pattern = r'\b(\d+)\s*(%s)\b' % '|'.join(re.escape(unit) for unit in units.keys())
        processed_line = re.sub(pattern, replace_unit, line)
        processed_lines.append(processed_line)
    return "\n".join(processed_lines)



def correct_scientific_units_with_logging(text):
    global global_logs
    unit_symbols = ['kg', 'm', 's', 'A', 'K', 'mol', 'cd', 'Hz', 'N', 'Pa', 'J', 'W', 'C', 'V', 'F', 'Ω', 'ohm', 'S', 'T', 'H', 'lm', 'lx', 'Bq', 'Gy', 'Sv', 'kat']
    pattern = rf"\b(\d+)\s*({'|'.join(re.escape(unit) for unit in unit_symbols)})\s*(s|'s|\.s)?\b"
    lines = text.splitlines()
    updated_lines = []
    
    for line_number, line in enumerate(lines, start=1):
        original_line = line
        changes = []
        new_line = re.sub(pattern, lambda m: f"{m.group(1)} {m.group(2)}", line)
                
        if new_line != line:
            for match in re.finditer(pattern, line):
                original = match.group(0)
                corrected = f"{match.group(1)} {match.group(2)}"
                if original != corrected:
                    changes.append(f"'{original}' -> '{corrected}'")

            if changes:
                global_logs.append(
                    f"[unit correction] Line {line_number}: {', '.join(changes)}"
                )

        updated_lines.append(new_line)
        
    return "\n".join(updated_lines)



def write_to_log(doc_id):
    global global_logs

    output_dir = os.path.join('output', str(doc_id))
    os.makedirs(output_dir, exist_ok=True)
    log_file_path = os.path.join(output_dir, 'global_logs.txt')

    with open(log_file_path, 'w', encoding='utf-8') as log_file:
        log_file.write("\n".join(global_logs))

    global_logs = []







# twofold not two-fold hyphenate with numeral for numbers greater than nine, e.g. 10-fold. 
def replace_fold_phrases(text):
    """
    Replaces phrases with '-fold' to ensure correct formatting based on the number preceding it.
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs
    def process_fold(match):
        original = match.group(0)
        num_str = match.group(1)
        separator = match.group(2)
        
        if separator != "-":
            return original

        try:
            if num_str.isdigit():
                number = int(num_str)
            else:
                number = w2n.word_to_num(num_str)

            if number > 9:
                modified = f"{number}-fold"
            else:
                modified = f"{num2words(number)}fold"

            if original != modified:
                line_number = text[:match.start()].count('\n') + 1
                global_logs.append(
                    f"[replace_fold_phrases] Line {line_number}: '{original}' -> '{modified}'"
                )
            return modified
        except ValueError:
            return original

    pattern = r"(\b\w+\b)(-?)fold"
    updated_text = re.sub(pattern, process_fold, text)
    return updated_text





# def correct_preposition_usage(text):
#     def process_from_to(match):
#         return f"from {match.group(1)} to {match.group(2)}"

#     def process_between_and(match):
#         return f"between {match.group(1)} and {match.group(2)}"
#     text = re.sub(r"from (\d{4})[–-](\d{4})", process_from_to, text)
#     text = re.sub(r"between (\d{4})[–-](\d{4})", process_between_and, text)
#     return text




def correct_preposition_usage(text):
    """
    Corrects preposition usage for date ranges (e.g., "from 2000-2010" -> "from 2000 to 2010").
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs
    def process_from_to(match):
        original = match.group(0)
        modified = f"from {match.group(1)} to {match.group(2)}"
        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[correct_preposition_usage] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified
    def process_between_and(match):
        original = match.group(0)
        modified = f"between {match.group(1)} and {match.group(2)}"
        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[correct_preposition_usage] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified
    text = re.sub(r"from (\d{4})[–-](\d{4})", process_from_to, text)
    text = re.sub(r"between (\d{4})[–-](\d{4})", process_between_and, text)
    return text





def correct_scientific_unit_symbols(text):
    """
    Ensures proper capitalization of units derived from proper names (e.g., J, Hz, W, N) only when preceded by a number.
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs
    units = {
        "j": "J",
        "hz": "Hz",
        "w": "W",
        "n": "N",
        "pa": "Pa",
        "v": "V",
        "a": "A",
        "c": "C",
        "lm": "lm",
        "lx": "lx",
        "t": "T",
        "ohm": "Ω",
        "s": "S",
        "k": "K",
        "cd": "cd",
        "mol": "mol",
        "rad": "rad",
        "sr": "sr"
    }
    def process_unit(match):
        original = match.group(0)
        unit = match.group(2).lower()
        modified = f"{match.group(1)}{units.get(unit, match.group(2))}"
        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[correct_scientific_unit_symbols] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified
    pattern = r"\b(\d+\s*)(%s)\b" % "|".join(re.escape(unit) for unit in units.keys())
    updated_text = re.sub(pattern, process_unit, text, flags=re.IGNORECASE)
    return updated_text




# def remove_quotation(text: str):
#     #  remove quotation '
#       para_text = re.sub(r"([A-Z]+)'", r'\1',text)
#       return para_text


def remove_quotation(text: str):
    """
    Removes single quotation marks (') following capitalized words.
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs
    pattern = r"([A-Z]+)'"
    def process_quotation_removal(match):
        original = match.group(0)
        modified = f"{match.group(1)}"

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[remove_quotation] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified
    para_text = re.sub(pattern, process_quotation_removal, text)
    return para_text






# def remove_and(text:str):
#     # Load the document
#     #doc = Document(file_path)
#     # Regex pattern to match "and" between two capitalized words
#     pattern = r'([A-Z][a-z]+)\s+and\s+([A-Z][a-z]+)'
#     text = re.sub(pattern, r'\1 & \2',text)
#     text = re.sub(pattern, r'\1 & \2',text)
#     return text



def remove_and(text: str):
    """
    Replaces 'and' between two capitalized words with an ampersand (&).
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs
    pattern = r'([A-Z][a-z]+)\s+and\s+([A-Z][a-z]+)'
    def process_and_replacement(match):
        original = match.group(0)
        modified = f"{match.group(1)} & {match.group(2)}"
        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[remove_and] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified
    text = re.sub(pattern, process_and_replacement, text)
    return text




def correct_units_in_ranges_with_logging(text):
    global global_logs

    # List of valid unit symbols
    unit_symbols = ['cm', 'm', 'kg', 's', 'A', 'K', 'mol', 'cd', '%']

    # Regex patterns
    range_pattern = rf"\b(\d+)\s*({'|'.join(re.escape(unit) for unit in unit_symbols)})\s*(to|-|–|—)\s*(\d+)\s*\2\b"
    thin_space_pattern = rf"\b(\d+)\s+({'|'.join(re.escape(unit) for unit in unit_symbols)})\b"

    lines = text.splitlines()
    updated_lines = []

    for line_number, line in enumerate(lines, start=1):
        original_line = line

        # Correct repeated units in ranges
        new_line = re.sub(
            range_pattern,
            lambda m: f"{m.group(1)} {m.group(3)} {m.group(4)} {m.group(2)}",
            line
        )

        # Add thin space between value and unit (except %)
        new_line = re.sub(
            thin_space_pattern,
            lambda m: f"{m.group(1)} {m.group(2)}" if m.group(2) != "%" else f"{m.group(1)}{m.group(2)}",
            new_line
        )

        # Log changes if any
        if new_line != line:
            change_details = f"{line.strip()} -> {new_line.strip()}"
            global_logs.append(f"Line {line_number}: {change_details}")
            line = new_line

        updated_lines.append(line)

    # Return the updated text
    return "\n".join(updated_lines)







def correct_unit_spacing(text):
    """
    Corrects spacing between numbers and units (e.g., "100 MB" -> "100MB").
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs
    units = ["Hz", "KHz", "MHz", "GHz", "kB", "MB", "GB", "TB"]
    pattern = r"(\d+)\s+(" + "|".join(units) + r")"
    def process_spacing(match):
        original = match.group(0)
        modified = f"{match.group(1)}{match.group(2)}"
        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[correct_unit_spacing] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified
    corrected_text = re.sub(pattern, process_spacing, text)
    return corrected_text






def apply_quotation_punctuation_rule(text: str):
    """
    Adjusts the placement of punctuation marks within single quotes.
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs
    pattern = r"‘(.*?)’([!?])"
    def process_quotation_punctuation(match):
        original = match.group(0)
        modified = f"‘{match.group(1)}{match.group(2)}’"
        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[apply_quotation_punctuation_rule] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified
    updated_text = re.sub(pattern, process_quotation_punctuation, text)
    return updated_text



# def enforce_dnase_rule(text: str):
#     pattern = r"\bDNAse\b"
#     updated_text = re.sub(pattern, "DNase", text)
#     return updated_text



def enforce_dnase_rule(text: str):
    """
    Enforces the correct capitalization for 'DNase'.
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs

    pattern = r"\bDNAse\b"

    def process_dnase_replacement(match):
        original = match.group(0)
        modified = "DNase"

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[enforce_dnase_rule] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified

    updated_text = re.sub(pattern, process_dnase_replacement, text)
    return updated_text



def apply_remove_italics_see_rule(text):
    return text.replace('*see*', 'see')



# There is one problem here for project, & document it is not changing and for project & document it is changing
def replace_ampersand(text):
    global global_logs
    def replacement(match):
        left, right = match.group(1), match.group(2)
        original = match.group(0)
        line_number = text[:match.start()].count('\n') + 1
        if left[0].isupper() and right[0].isupper():
            return original

        modified = left + ' and ' + right
        global_logs.append(
            f"[replace_ampersand] Line {line_number}: '{original}' -> '{modified}'"
        )
        return modified
    return re.sub(r'(?m)(\w+)\s*&\s*(\w+)', replacement, text)


def rename_section(text):
    # Replace all occurrences of the § symbol with 'Section'
    return re.sub(r'§', 'Section', text)




def process_url_add_http(text):
    """
    Adjusts URLs in the input text based on the given rules:
    1. If a URL starts with 'www.' but doesn't have 'http://', prepend 'http://'.
    2. If a URL already starts with 'http://', remove 'http://'.

    Args:
        text (str): The input text containing URLs.

    Returns:
        str: The modified text with URLs adjusted.
    """
    global global_logs

    def add_http_prefix(match):
        original = match.group(0)
        modified = f"http://{match.group(1)}"

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[process_url_add_http] Line {line_number}: '{original}' -> '{modified}'"
            )

        return modified

    def remove_http_prefix(match):
        original = match.group(0)
        modified = match.group(1)

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1  # Calculate the line number
            global_logs.append(
                f"[process_url_add_http] Line {line_number}: '{original}' -> '{modified}'"
            )

        return modified

    text = re.sub(r"\bhttp://(www\.\S+)", remove_http_prefix, text)
    text = re.sub(r"\b(www\.\S+)", add_http_prefix, text)
    return text





def process_url_remove_http(url):
    """
    Removes 'http://' from a URL if there is no path, parameters, query, or fragment.
    Args:
        url (str): The input URL to process.
    Returns:
        str: The modified URL with 'http://' removed if applicable.
    """
    global global_logs

    parsed = urlparse(url)
    original = url

    if parsed.scheme == "http" and not (parsed.path or parsed.params or parsed.query or parsed.fragment):
        modified = parsed.netloc

        if original != modified:
            line_number = 1
            global_logs.append(
                f"[process_url_remove_http] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified
    return url



def process_symbols_mark(text, line_number, symbols=["®", "™", "©", "℗", "℠"]):
    """
    Ensures symbols like ®, ™, etc., appear only the first time in the text.
    Updates the global_log with changes, including line number, original text, and updated text.
    """
    original_text = text
    symbol_set = set()
    global global_logs
    
    for symbol in symbols:
        occurrences = list(re.finditer(re.escape(symbol), text))
        if occurrences:
            first_occurrence = occurrences[0].start()
            # Replace all occurrences after the first one
            text = (
                text[:first_occurrence + 1]
                + re.sub(re.escape(symbol), "", text[first_occurrence + 1:])
            )
            symbol_set.add(symbol)

    # Log changes if the text was modified
    if original_text != text:
        global_logs.append(
            f"[process_symbols_in_doc] Line {line_number}: '{original_text}' -> '{text}'"
        )

    return text




def remove_commas_from_numbers(text, line_number):
    """
    Removes commas from numerical values in the text.
    Updates the global_log with the specific changes, including line number and changes made.
    """
    original_text = text
    changes = []
    global global_logs

    # Regex to match numbers with commas (e.g., 1,000 or 20,000)
    pattern = r'\b\d{1,3}(,\d{3})+\b'
    def replacer(match):
        original_number = match.group(0)  # Match the original number
        updated_number = original_number.replace(",", " ")  # Remove commas
        changes.append((original_number, updated_number))  # Log the change
        return updated_number
    text = re.sub(pattern, replacer, text)
    # Log individual changes
    for original, updated in changes:
        global_logs.append(
            f"[process_symbols_in_doc] Line {line_number}: '{original}' -> '{updated}'"
        )
    return text




def remove_spaces_from_four_digit_numbers(text, line_number):
    """
    Removes spaces from four-digit numerals in the text.
    Updates the global_log with specific changes, including line number and changes made.
    """
    original_text = text
    changes = []
    global global_logs

    pattern = r'\b\d\s\d{3}\b'

    def replacer(match):
        original_number = match.group(0)  # Match the original number
        updated_number = original_number.replace(" ", "")  # Remove spaces
        changes.append((original_number, updated_number))  # Log the change
        return updated_number

    text = re.sub(pattern, replacer, text)

    for original, updated in changes:
        global_logs.append(
            f"[process_symbols_in_doc] Line {line_number}: '{original}' -> '{updated}'"
        )

    return text



def set_latinisms_to_roman_in_runs(paragraph_text, line_number, latinisms=None):
    """
    Converts specific Latinisms from italic to roman text in a string of text.
    Logs changes to the global_log, including line number and original italicized Latinism.
    """
    if latinisms is None:
        latinisms = [
            "i.e.", "e.g.", "via", "vice versa", "etc.", "a posteriori", 
            "a priori", "et al.", "cf.", "c."
        ]
    changes = []
    global global_logs
    # Process the text, and for each Latinism, replace its italics if needed
    for lat in latinisms:
        if lat in paragraph_text:
            changes.append(lat)  # Log the Latinism that was changed

    # for changed in changes:
    #     global_logs.append(
    #         f"[process_symbols_in_doc] Line {line_number}: '{changed}' -> '{changed}'"
    #     )
    return paragraph_text 



def convert_decimal_to_baseline(paragraph_text, line_number):
    """
    Converts any non-standard decimal separator (•) to a standard decimal point (.)
    only when both sides are numeric.
    Logs the changes to the global_log, including line number and the change.
    """
    changes = []
    global global_logs
    # Regular expression to find '•' between numbers
    pattern = r'(?<=\d)\xB7(?=\d)'

    # Find all occurrences of '•' that are between digits and replace with '.'
    matches = re.findall(pattern, paragraph_text)
    if matches:
        original_text = paragraph_text
        updated_text = re.sub(pattern, '.', paragraph_text)  # Replace '•' with '.'
        changes.append((original_text, updated_text))

    for original, updated in changes:
        global_logs.append(
            f"[convert_decimal_to_baseline] Line {line_number}: '{original}' -> '{updated}'"
        )
    return updated_text if changes else paragraph_text



# Function to convert numbers to words (1 to 10)
def number_to_word(num):
    num_dict = {
        1: 'One', 2: 'Two', 3: 'Three', 4: 'Four', 5: 'Five',
        6: 'Six', 7: 'Seven', 8: 'Eight', 9: 'Nine', 10: 'Ten'
    }
    return num_dict.get(num, str(num))


# Function to convert words to numbers
def word_to_number(word):
    word_dict = {
        'one': 1, 'two': 2, 'three': 3, 'four': 4, 'five': 5,
        'six':6, 'seven':7, 'eight':8, 'nine':9, 'ten':10 
    }
    return word_dict.get(word.lower(), word)


# Function to process text and replace words with numbers, and numbers with words
def convert_text(text):
    text = re.sub(r'\b([1-9]|10)\b', lambda match: number_to_word(int(match.group(0))), text)
    text = re.sub(r'\b(one|two|three|four|five|six|seven|eight|nine|ten)\s*(kg|m|cm|g|l)\b', 
    lambda match: str(word_to_number(match.group(1))) + ' ' + match.group(2), text, flags=re.IGNORECASE)
    return text


def adjust_punctuation_style_using_paragraph_text(text, para_runs):
    """
    Analyze `text` to detect italicized or bold characters followed by punctuation
    and ensure the punctuation inherits the appropriate style (italic or bold).
    """
    for i in range(len(para_runs) - 1):
        current_run = para_runs[i]
        next_run = para_runs[i + 1]

        # Check if current run ends with italicized text
        if current_run.text and current_run.italic:
            last_char = current_run.text[-1]
            if next_run.text and next_run.text[0] in ".,!?\"'()":
                next_run.italic = True
        
        elif current_run.text and current_run.bold:
            last_char = current_run.text[-1]
            if next_run.text and next_run.text[0] in ".,!?\"”'()":
                next_run.bold = True
    
    # Return updated text after style adjustments
    return text

# Dictionary to convert word numbers to integer values
word_to_num = {
    'zero': 0, 'one': 1, 'two': 2, 'three': 3, 'four': 4, 'five': 5,
    'six': 6, 'seven': 7, 'eight': 8, 'nine': 9, 'ten': 10, 'eleven': 11,
    'twelve': 12, 'thirteen': 13, 'fourteen': 14, 'fifteen': 15, 'sixteen': 16,
    'seventeen': 17, 'eighteen': 18, 'nineteen': 19, 'twenty': 20, 'twenty-one': 21,
    'twenty-two': 22, 'twenty-three': 23, 'twenty-four': 24, 'twenty-five': 25,
    'twenty-six': 26, 'twenty-seven': 27, 'twenty-eight': 28, 'twenty-nine': 29,
    'thirty': 30, 'thirty-one': 31, 'thirty-two': 32, 'thirty-three': 33,
    'thirty-four': 34, 'thirty-five': 35, 'thirty-six': 36, 'thirty-seven': 37,
    'thirty-eight': 38, 'thirty-nine': 39, 'forty': 40, 'forty-one': 41,
    'forty-two': 42, 'forty-three': 43, 'forty-four': 44, 'forty-five': 45,
    'forty-six': 46, 'forty-seven': 47, 'forty-eight': 48, 'forty-nine': 49,
    'fifty': 50, 'fifty-one': 51, 'fifty-two': 52, 'fifty-three': 53,
    'fifty-four': 54, 'fifty-five': 55, 'fifty-six': 56, 'fifty-seven': 57,
    'fifty-eight': 58, 'fifty-nine': 59, 'sixty': 60, 'sixty-one': 61,
    'sixty-two': 62, 'sixty-three': 63, 'sixty-four': 64, 'sixty-five': 65,
    'sixty-six': 66, 'sixty-seven': 67, 'sixty-eight': 68, 'sixty-nine': 69,
    'seventy': 70, 'seventy-one': 71, 'seventy-two': 72, 'seventy-three': 73,
    'seventy-four': 74, 'seventy-five': 75, 'seventy-six': 76, 'seventy-seven': 77,
    'seventy-eight': 78, 'seventy-nine': 79, 'eighty': 80, 'eighty-one': 81,
    'eighty-two': 82, 'eighty-three': 83, 'eighty-four': 84, 'eighty-five': 85,
    'eighty-six': 86, 'eighty-seven': 87, 'eighty-eight': 88, 'eighty-nine': 89,
    'ninety': 90, 'ninety-one': 91, 'ninety-two': 92, 'ninety-three': 93,
    'ninety-four': 94, 'ninety-five': 95, 'ninety-six': 96, 'ninety-seven': 97,
    'ninety-eight': 98, 'ninety-nine': 99, 'hundred': 100
}

# Reverse dictionary to convert integer values back to words
num_to_word = {v: k for k, v in word_to_num.items()}

# Function to convert word number to integer
def word_to_int(word):
    return word_to_num.get(word.lower(), None)

# Function to convert integer to word
def int_to_word(num):
    return num_to_word.get(num, None)

# Regular expression to match "word and word" pattern
pattern = re.compile(r'(\b\w+\b) and (\b\w+\b)')

# Function to process the string with regex and apply transformations
def process_string(text):
    def replace_match(match):
        word1 = match.group(1)
        word2 = match.group(2)
        # Convert words to their numeric values
        num1 = word_to_int(word1)
        num2 = word_to_int(word2)
        
        # If both numbers are less than 9, return them as word form
        if (num1 is not None and num1 < 9) and (num2 is not None and num2 < 9):
            return f"{word1} and {word2}"  # No change if both are < 9
        
        # If either number is greater than or equal to 9, convert to numeric form
        if (num1 is not None and num1 >= 9) or (num2 is not None and num2 >= 9):
            # Convert both to numeric form
            num1 = num1 if num1 is not None else word1
            num2 = num2 if num2 is not None else word2
            return f"{num1} and {num2}"  # Replace with numeric values
        
        return match.group(0)  # Return the match as is if both are < 9
    
    # Apply regex substitution with the replace function
    return pattern.sub(replace_match, text)
 

 
# put space between 
def format_hyphen_to_en_dash(runs, line_number):
    """
    Replace hyphens with en dashes in a Word document paragraph's runs.
    Adjust spacing based on surrounding context:
    - Add spaces if there are words on both sides.
    - Remove spaces if there are numbers on both sides.
    Logs changes to the global 'global_logs' list.
    :param runs: The runs of a paragraph (doc.paragraphs[n].runs)
    :param line_number: The line number of the paragraph being processed
    """
    word_range_pattern = re.compile(r'(\b\w+)\s*-\s*(\w+\b)')
    number_range_pattern = re.compile(r'(\d+)\s*-\s*(\d+)')
    for run in runs:
        if run.text:
            original_text = run.text
            # Replace hyphen with en dash and remove spaces for number ranges
            updated_text = number_range_pattern.sub(r'\1–\2', original_text)
            # Replace hyphen with en dash and ensure spaces for word ranges
            updated_text = word_range_pattern.sub(r'\1 – \2', updated_text)
            if updated_text != original_text:
                # Log the change
                global_logs.append(
                    f"Line {line_number}: '{original_text}' -> '{updated_text}'"
                )
                # Update the run text
                run.text = updated_text


def replace_em_with_en(runs, line_number):
    """
    Replaces all em dashes (—) with en dashes (–) in the text of a paragraph's runs.    
    Args:
        runs: The runs of a paragraph (e.g., `para.runs`).
        line_number: The line number of the paragraph for context (not used in the function directly).
    """
    for run in runs:
        if '—' in run.text:
            run.text = run.text.replace('—', '–')




def replace_dashes(runs, line_number):
    """
    Replaces em dashes (—) and normal hyphens (-) with en dashes (–) in the text of a paragraph's runs.
    Logs changes to a global list with details of the modification in the desired format.
    Args:
        runs: The runs of a paragraph (e.g., `para.runs`).
        line_number: The line number of the paragraph for context.
    """
    global global_logs
    for run in runs:
        original_text = run.text
        modified_text = run.text.replace('—', '–').replace('-', '–')
        if original_text != modified_text:
            run.text = modified_text
            global_logs.append(
                f"[replace_dashes_with_logging] Line {line_number}: '{original_text}' -> '{modified_text}'"
            )



def convert_currency_to_symbols(runs, line_number):
    """
    Converts textual currency names (dollar, pound, euro) to symbols ($, £, €) 
    when preceded by a numerical value in the text of a paragraph's runs.
    Logs changes to a global list with details of the modification in the desired format.    
    Args:
        runs: The runs of a paragraph (e.g., `para.runs`).
        line_number: The line number of the paragraph for context.
    """
    global global_logs
    currency_patterns = {
        r'(\b\d+\s*)dollars\b': r'$\1',
        r'(\b\d+\s*)pounds\b': r'£\1',
        r'(\b\d+\s*)euros\b': r'€\1'
    }

    for run in runs:
        original_text = run.text
        modified_text = original_text
        # Apply each currency replacement pattern
        for pattern, replacement in currency_patterns.items():
            modified_text = re.sub(pattern, replacement, modified_text, flags=re.IGNORECASE)
        # If changes are made, update the text and log the change
        if original_text != modified_text:
            run.text = modified_text
            global_logs.append(
                f"[convert_currency_to_symbols] Line {line_number}: '{original_text}' -> '{modified_text}'"
            )



def curly_to_straight(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = run.text.replace('“', '"').replace('”', '"').replace('‘', "'").replace('’', "'")
            



# def staright_to_curly(doc):
#     for para in doc.paragraphs:
#         para.text = replace_straight_quotes_with_curly(para.text)

def straight_to_curly(doc):
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = replace_straight_quotes_with_curly(run.text)
        # Rebuild paragraph text, preserving spacing
        paragraph_text = " ".join([run.text for run in para.runs])
        para.clear()  # Clear current paragraph content
        para_runs = re.sub(r'\s([.,])',r'\1',paragraph_text)   
        para.add_run(para_runs)  # Re-add the updated text into the paragraphgraph_text)  # Re-add the updated text into the paragraph



# def highlight_and_correct(doc):
#     """
#     This function highlights incorrectly spelled words in a Word document by changing their font color to red.
#     Words enclosed in single or double quotes are ignored.
#     Args:
#         doc: The Word document object (from python-docx).
#         us_dict: A spell-checking dictionary object (e.g., from the `pyspellchecker` library).
#     """
#     for para in doc.paragraphs:
#         formatted_runs = []

#         for run in para.runs:
#             words = run.text.split()
#             for i, word in enumerate(words):
#                 original_word = word 
#                 punctuation = ""

#                 # Separate trailing punctuation (if any)
#                 if word[-1] in ",.?!:;\"'()[]{}":
#                     punctuation = word[-1]
#                     word = word[:-1]

#                 # Ignore words fully enclosed in single or double quotes
#                 if (word.startswith('"') and word.endswith('"')) or (word.startswith("'") and word.endswith("'")):
#                     formatted_runs.append((original_word, None))
#                 # Ignore empty words
#                 elif not word.strip():
#                     formatted_runs.append((original_word, None))
#                 # Check spelling and mark incorrect words in red
#                 elif not us_dict.check(word.lower()):
#                     formatted_runs.append((word, RGBColor(255, 0, 0)))  # Highlight misspelled word
#                 else:
#                     formatted_runs.append((word, None))  # Correct word

#                 # Add punctuation back to the word, if it had any
#                 if punctuation:
#                     formatted_runs.append((punctuation, None))

#                 # Add a space after the word unless it's the last one
#                 if i < len(words) - 1:
#                     formatted_runs.append((" ", None))

#         # Clear the paragraph's text and rebuild it with formatted runs
#         para.clear()  # Clear existing paragraph content

#         for text, color in formatted_runs:
#             new_run = para.add_run(text)  # Add new text to the paragraph
#             if color:  # If a color is specified, apply it
#                 new_run.font.color.rgb = color


def highlight_and_correct(doc):
    """
    Highlights incorrectly spelled words in a Word document by changing their font color to red.
    Formatting of the document remains unchanged.
    Args:
        doc: The Word document object (from python-docx).
        us_dict: A spell-checking dictionary object (e.g., from the `pyspellchecker` library).
    """
    for para in doc.paragraphs:
        for run in para.runs:
            words = run.text.split()
            updated_text = []
            for word in words:
                original_word = word
                punctuation = ""

                # Separate trailing punctuation (if any)
                if word and word[-1] in ",.?!:;\"'()[]{}":
                    punctuation = word[-1]
                    word = word[:-1]

                # Ignore words fully enclosed in single or double quotes
                if (word.startswith('"') and word.endswith('"')) or (word.startswith("'") and word.endswith("'")):
                    updated_text.append(original_word)
                # Ignore empty words
                elif not word.strip():
                    updated_text.append(original_word)
                # Check spelling and mark incorrect words in red
                elif not us_dict.check(word.lower()):
                    updated_text.append(word)
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Highlight misspelled word
                else:
                    updated_text.append(word)  # Correct word remains unchanged

                # Add punctuation back to the word, if it had any
                if punctuation:
                    updated_text.append(punctuation)

            # Preserve formatting by updating text in place
            run.text = " ".join(updated_text)


def clean_word1(word):
    return ''.join(filter(str.isalnum, word)).lower()


# Helper function to extract text from docx file
def extract_text_from_docx(file_path):
    try:
        with open(file_path, "rb") as docx_file:
            result = mammoth.extract_raw_text(docx_file)
            return result.value
    except Exception as e:
        # logging.error(f"Error extracting text from file: {e}")
        return ""
    


SECRET_KEY = "Naveen"
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 600


class TokenRequest(BaseModel):
    token: str


@router.post("/process_us")
async def process_file(token_request: TokenRequest, doc_id: int = Query(...)):
    try:
        payload = jwt.decode(token_request.token, SECRET_KEY, algorithms=[ALGORITHM])
        print("Decoded Token Data:", payload)
        # global global_logs
        conn = get_db_connection()
        if conn is None:
            raise HTTPException(status_code=500, detail="Database connection error")
        
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM row_document WHERE row_doc_id = %s", (doc_id,))
        rows = cursor.fetchone()

        if not rows:
            raise HTTPException(status_code=404, detail="Document not found")
        
        file_path = os.path.join(os.getcwd(), 'files', rows[1])

        if not os.path.exists(file_path):
            raise HTTPException(status_code=404, detail="File not found on server")

        start_time = datetime.now()

        file_content = extract_text_from_docx(file_path)
        text = file_content

        global_logs.append(f"FileName: {rows[1]}\n\n")

        lines = text.split('\n')
        for index, line in enumerate(lines):
            words = line.split()
            for word in words:
                cleaned = clean_word1(word)
                if cleaned and not us_dict.check(cleaned):
                    suggestions = us_dict.suggest(cleaned)
                    suggestion_text = (
                        f" Suggestions: {', '.join(suggestions)}"
                        if suggestions else " No suggestions available"
                    )
                    global_logs.append(f"Line {index}: {word} ->{suggestion_text}")

        end_time = datetime.now()
        time_taken = round((end_time - start_time).total_seconds(), 2)
        time_log = f"\nStart Time: {start_time}\nEnd Time: {end_time}\nAnalysis completed in {time_taken} seconds.\n\n"

        global_logs.insert(0, time_log)

        document_name = rows[1].replace('.docx', '')
        log_filename = f"log_main.txt"
        
        output_path_file = Path(os.getcwd()) / 'output' / str(doc_id) / log_filename
        dir_path = output_path_file.parent

        dir_path.mkdir(parents=True, exist_ok=True)
        
        output_dir = os.path.join("output", str(doc_id))
        os.makedirs(output_dir, exist_ok=True)

        output_path = os.path.join(output_dir, f"processed_{os.path.basename(file_path)}")

        doc = docx.Document(file_path)

        curly_to_straight(doc)
        highlight_and_correct(doc)
        write_to_log(doc_id)
        
        process_doc_function1(payload, doc, doc_id)
        process_doc_function2(payload, doc, doc_id)
        process_doc_function3(payload, doc, doc_id)
        process_doc_function4(payload, doc, doc_id)
        process_doc_function6(payload, doc, doc_id)
        process_doc_function7(payload, doc, doc_id)
         
        straight_to_curly(doc)
        
        doc.save(output_path)

        cursor.execute("SELECT final_doc_id FROM final_document WHERE row_doc_id = %s", (doc_id,))
        existing_rows = cursor.fetchall()

        if existing_rows:
            logging.info('File already processed in final_document. Skipping insert.')
        else:
            folder_url = f'/output/{doc_id}/'
            cursor.execute(
                '''INSERT INTO final_document (row_doc_id, user_id, final_doc_size, final_doc_url, status, creation_date)
                VALUES (%s, %s, %s, %s, %s, NOW())''',
                (doc_id, rows[1], rows[2], folder_url, rows[7])
            )
            logging.info('New file processed and inserted into final_document.')

        conn.commit()
        # write_to_log(doc_id)
        logging.info(f"Processed file stored at: {output_path}")
        return {"success": True, "message": f"File processed and stored at {output_path}"}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



class CheckboxData(RootModel[Dict[int, bool]]):
    """Root model to accept numeric keys with boolean values."""

class TokenResponse(BaseModel):
    token: str

class TokenRequest(BaseModel):
    token: str


    
    
@router.post("/generate-token", response_model=TokenResponse)
async def generate_token(checkbox_data: CheckboxData):
    """
    API endpoint to generate a JWT token from checkbox data with numeric keys.
    """
    try:
        data = checkbox_data.root
        to_encode = data.copy()

        expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        to_encode.update({"exp": expire})

        token = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
        return {"token": token}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating token: {str(e)}")



# Decode and Use Token API
@router.post("/use-token")
async def use_token(token_request: TokenRequest):
    try:
        # Decode the token
        payload = jwt.decode(token_request.token, SECRET_KEY, algorithms=[ALGORITHM])
        return {"message": "Token processed successfully!", "data": payload}
    except JWTError:
        raise HTTPException(status_code=401, detail="Invalid or expired token")












@router.get("/rules", summary="Get all rules", response_description="List of rules")
def get_rules():
    conn = get_db_connection()
    try:
        # Fetch all rules
        cursor = conn.cursor()
        cursor.execute("SELECT id, rule_name FROM rules")
        rows = cursor.fetchall()
        print(rows)
        
        if not rows:
            raise HTTPException(status_code=404, detail="No rules found")
        
        # Convert rows to a list of dictionaries
        rules = [{"id": row[0], "rule_name": row[1]} for row in rows]
        return {"success": True, "data": rules}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        conn.close()
