import re
from docx import Document
import os
from urllib.parse import urlparse
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from pathlib import Path
from datetime import datetime

# Global logs to keep track of changes
global_logs = []


def clean_web_addresses(runs):
    """
    Removes angle brackets around web addresses (e.g., "<http://example.com>" -> "http://example.com").
    Args:
        runs (list): A list of runs (segments of text in the document).
    """
    def process_web_address(match, text):
        original = match.group(0)
        modified = match.group(1)

        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[clean_web_addresses] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified

    for run in runs:
        # Apply cleaning to each run's text in place
        run.text = re.sub(r"<(https?://[^\s<>]+)>", lambda match: process_web_address(match, run.text), run.text)


def remove_concluding_slashes_from_urls(document):
    """Update hyperlink targets and visible text (including in tables)."""
    rels = document.part.rels
    hyperlink_rels = {rel_id: rel for rel_id, rel in rels.items() if rel.reltype == RT.HYPERLINK}

    global global_logs
    for rel_id, rel in hyperlink_rels.items():
        original_url = rel._target
        if not original_url.endswith('/'):
            continue

        updated_url = original_url.rstrip('/')
        rel._target = updated_url  # Update target URL
        global_logs.append(f"Updated target: {original_url} -> {updated_url}")

        # Update visible text in paragraphs and tables
        for element in document.element.body:
            # Process paragraphs
            if element.tag.endswith('p'):
                for run in element.xpath('.//w:r'):
                    text = run.xpath('.//w:t')
                    if text and text[0].text == original_url:
                        text[0].text = updated_url
                        global_logs.append(f"Updated paragraph text: {original_url} -> {updated_url}")

            # Process tables
            elif element.tag.endswith('tbl'):
                for cell in element.xpath('.//w:tc'):
                    for run in cell.xpath('.//w:r'):
                        text = run.xpath('.//w:t')
                        if text and text[0].text == original_url:
                            text[0].text = updated_url
                            global_logs.append(f"Updated table text: {original_url} -> {updated_url}")


def drop_https(document):
    """
    Update hyperlink targets and visible text (including in tables) according to the rule:
    - If the URL points only to a domain (i.e. path is '' or '/' with no query or fragment),
      drop the scheme (http:// or https://) so that only the domain remains.
    - Otherwise, leave the URL unchanged.
    """
    # Get all hyperlink relationships from the document
    rels = document.part.rels
    hyperlink_rels = {rel_id: rel for rel_id, rel in rels.items() if rel.reltype == RT.HYPERLINK}

    global global_logs
    for rel_id, rel in hyperlink_rels.items():
        original_url = rel._target  # using the internal attribute to access the target
        parsed = urlparse(original_url)

        # Determine if URL is a bare domain (no path beyond "/" and no query/fragment)
        if parsed.path in ['', '/'] and not parsed.query and not parsed.fragment:
            new_url = parsed.netloc  # drop scheme, e.g. "http://example.com" -> "example.com"
        else:
            new_url = original_url

        # Only update if the URL has changed
        if new_url != original_url:
            # Update the hyperlink relationship target
            rel._target = new_url
            global_logs.append(f"Updated target: {original_url} -> {new_url}")

            # Update visible text in paragraphs and tables if it exactly matches the original URL
            for element in document.element.body:
                # Process paragraphs
                if element.tag.endswith('p'):
                    for run in element.xpath('.//w:r'):
                        text_elems = run.xpath('.//w:t')
                        if text_elems and text_elems[0].text == original_url:
                            text_elems[0].text = new_url
                            global_logs.append(f"Updated paragraph text: {original_url} -> {new_url}")
                # Process tables
                elif element.tag.endswith('tbl'):
                    for cell in element.xpath('.//w:tc'):
                        for run in cell.xpath('.//w:r'):
                            text_elems = run.xpath('.//w:t')
                            if text_elems and text_elems[0].text == original_url:
                                text_elems[0].text = new_url
                                global_logs.append(f"Updated table text: {original_url} -> {new_url}")



def process_url_remove_http(runs):
    """
    Removes 'http://' from a URL if there is no path, parameters, query, or fragment.
    Args:
        runs (list): A list of runs (segments of text in the document).
    """
    def process_url(match, text):
        original = match.group(0)
        parsed = urlparse(original)
        
        # Remove 'http://' if there is no path, parameters, query, or fragment
        if parsed.scheme == "http" and not (parsed.path or parsed.params or parsed.query or parsed.fragment):
            modified = parsed.netloc
            if original != modified:
                line_number = text[:match.start()].count('\n') + 1
                global_logs.append(
                    f"[process_url_remove_http] Line {line_number}: '{original}' -> '{modified}'"
                )
            return modified
        return original

    for run in runs:
        # Apply the changes in place to each run's text
        run.text = re.sub(r"\bhttp://[^\s]+", lambda match: process_url(match, run.text), run.text)





def remove_hyperlinks_underline(document):
    # Process main document paragraphs
    for para in document.paragraphs:
        for hyperlink in para.hyperlinks:
            for run in hyperlink.runs:
                run.font.underline = False
    
    # Process tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for hyperlink in para.hyperlinks:
                        for run in hyperlink.runs:
                            run.font.underline = False
    
    # Process headers and footers
    for section in document.sections:
        # Process headers
        header = section.header
        for para in header.paragraphs:
            for hyperlink in para.hyperlinks:
                for run in hyperlink.runs:
                    run.font.underline = False
        # Process footers
        footer = section.footer
        for para in footer.paragraphs:
            for hyperlink in para.hyperlinks:
                for run in hyperlink.runs:
                    run.font.underline = False





def format_urls_in_paragraph(para):
    """
    Format URLs in paragraph runs to match the paragraph's base font.
    Handles both plain-text URLs and hyperlink URLs.
    """
    # Regex pattern for URL detection (improved)
    url_pattern = re.compile(
        r'(https?://|www\.)[\w\-\.]+\.[a-zA-Z]{2,}(/\S*)?',
        re.IGNORECASE
    )

    # Get paragraph's base font properties
    base_font = {
        'name': para.style.font.name,
        'size': para.style.font.size,
        'bold': para.style.font.bold,
        'italic': para.style.font.italic,
        'color': para.style.font.color.rgb if para.style.font.color else None
    }

    # Process plain-text URLs in runs
    full_text = ''.join(run.text for run in para.runs)
    url_matches = list(url_pattern.finditer(full_text))

    if url_matches:
        current_pos = 0
        for run in para.runs:
            run_text = run.text
            run_length = len(run_text)
            run_start = current_pos
            run_end = current_pos + run_length

            # Check if this run overlaps with any URL match
            for match in url_matches:
                url_start, url_end = match.start(), match.end()
                if run_start < url_end and url_start < run_end:
                    # Apply base font formatting to the entire run
                    run.font.name = base_font['name']
                    if base_font['size']:
                        run.font.size = base_font['size']
                    run.font.bold = base_font['bold']
                    run.font.italic = base_font['italic']
                    if base_font['color']:
                        run.font.color.rgb = base_font['color']
                    break  # Format once per run

            current_pos = run_end

    # Process hyperlinks (ensure display text matches base font)
    hyperlinks = para.hyperlinks
    for hyperlink in hyperlinks:
        if hyperlink.text:
            # Format hyperlink display text
            for run in hyperlink.runs:
                run.font.name = base_font['name']
                if base_font['size']:
                    run.font.size = base_font['size']
                run.font.bold = base_font['bold']
                run.font.italic = base_font['italic']
                if base_font['color']:
                    run.font.color.rgb = base_font['color']


def add_http_to_urls(document):
    pattern = re.compile(r'(?<!http://)(?<!https://)(www\.[^\s]+)')
    
    for para in document.paragraphs:
        # Assuming each paragraph has a 'hyperlinks' property containing hyperlink objects
        for hyperlink in para.hyperlinks:
            for run in hyperlink.runs:
                # Replace occurrences of 'www.' URLs without scheme with the proper prefix.
                run.text = pattern.sub(r'http://\1', run.text)

def write_to_log(doc_id, user):
    global global_logs
    current_date = datetime.now().strftime("%Y-%m-%d")
    output_path_file = Path(os.getcwd()) / 'output' / user / current_date / str(doc_id) / 'text' 
    os.makedirs(output_path_file, exist_ok=True)
    log_file_path = os.path.join(output_path_file, 'global_logs.txt')

    with open(log_file_path, 'a', encoding='utf-8') as log_file:
        log_file.write("\n".join(global_logs))

    global_logs = []
    


def process_doc_function4(payload: dict, doc: Document, doc_id,user):
    """
    This function processes the document by converting century notations
    and highlighting specific words.
    """
    line_number = 1
    # drop_https(doc)
    remove_concluding_slashes_from_urls(doc)
    remove_hyperlinks_underline(doc)
    # add_http_to_urls(doc)
    for para in doc.paragraphs:
    #     format_urls_in_paragraph(para)
        clean_web_addresses(para.runs)

       
    write_to_log(doc_id,user)
