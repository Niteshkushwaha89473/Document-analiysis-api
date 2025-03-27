import re
from docx import Document
import os
import roman
from word2number import w2n
import zipfile
import xml.etree.ElementTree as ET
from docx import Document
from db_config import get_db_connection
from fastapi import APIRouter, HTTPException, Depends, Query
from datetime import datetime
from pathlib import Path

global_logs = []

# def apply_sentence_case_footnotes(runs, line_number):
#     """
#     Converts the text in the given runs to sentence case.
#     - The first letter of the first run is capitalized.
#     - The rest of the text is converted to lowercase (except existing uppercase letters in abbreviations, etc.).
#     """
#     if not runs:
#         return
    
#     text = "".join(run.text for run in runs)
#     if not text.strip():
#         return
    
#     sentence_cased_text = text[0].upper() + text[1:].lower() if len(text) > 1 else text.upper()
    
#     start = 0
#     for run in runs:
#         run_length = len(run.text)
#         run.text = sentence_cased_text[start:start + run_length]
#         start += run_length




def extract_footnotes(docx_path):
    """
    Extracts footnotes from a DOCX file.
    Returns a list of tuples: (footnote_id, footnote_text).
    """
    footnotes = []
    # Open the docx file as a zip archive
    with zipfile.ZipFile(docx_path, 'r') as z:
        # Check if footnotes.xml exists in the archive
        if 'word/footnotes.xml' in z.namelist():
            with z.open('word/footnotes.xml') as f:
                tree = ET.parse(f)
                root = tree.getroot()
                # Define the XML namespace
                ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                # Iterate through each footnote element
                # print(dir(ns))
                for footnote in root.findall('w:footnote', ns):
                    footnote_id = footnote.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    if int(footnote_id) < 1:
                        continue
                    texts = []
                    # Each footnote can have several paragraphs
                    for p in footnote.findall('w:p', ns):
                        para_text = ""
                        # Extract text runs within the paragraph
                        for r in p.findall('w:r', ns):
                            for t in r.findall('w:t', ns):
                                para_text += t.text if t.text else ""
                                # print(para_text)
                        if para_text:
                            texts.append(para_text)
                    # Combine paragraphs into one string (separated by newlines)
                    combined_text = "\n".join(texts)
                    # print(combined_text)
                    footnotes.append((footnote_id, combined_text))
            return footnotes
        else:
            print("No footnotes.xml found in the document.")
            return []



def write_footnotes_to_docx(footnotes, output_path):
    """
    Writes the extracted footnotes to a new DOCX file.
    Each footnote is written as a new section with a heading.
    """
    doc = Document()
    doc.add_heading('Extracted Footnotes', level=1)
    for footnote_id, text in footnotes:
        if footnote_id=="1":
            continue
        doc.add_heading(f'Footnote {int(footnote_id)-1}', level=2)
        doc.add_paragraph(text)
    doc.save(output_path)




def convert_footnotes_to_endnotes(input_docx, output_docx):
    """
    Converts footnotes in the input DOCX to endnotes and writes the result to output_docx.
    This includes:
      - Replacing <w:footnoteReference> with <w:endnoteReference> in document.xml.
      - Converting word/footnotes.xml to word/endnotes.xml (changing tags appropriately).
      - Updating [Content_Types].xml: removing the footnotes override and adding an endnotes override.
      - Updating word/_rels/document.xml.rels: changing the relationship for footnotes to endnotes.
    """
    # Namespaces
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    
    updated_files = {}
    
    with zipfile.ZipFile(input_docx, 'r') as zin:
        for file in zin.namelist():
            data = zin.read(file)
            
            if file == "word/document.xml":
                tree = ET.ElementTree(ET.fromstring(data))
                root = tree.getroot()
                for elem in root.iter():
                    if elem.tag == f"{{{w_ns}}}footnoteReference":
                        elem.tag = f"{{{w_ns}}}endnoteReference"
                updated_data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                updated_files[file] = updated_data

            elif file == "word/footnotes.xml":
                tree = ET.ElementTree(ET.fromstring(data))
                root = tree.getroot()
                root.tag = f"{{{w_ns}}}endnotes"
                for child in root:
                    if child.tag == f"{{{w_ns}}}footnote":
                        child.tag = f"{{{w_ns}}}endnote"
                new_data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                updated_files["word/endnotes.xml"] = new_data
                print(" ")
                
            elif file == "[Content_Types].xml":
                tree = ET.ElementTree(ET.fromstring(data))
                root = tree.getroot()
                overrides_to_remove = []
                for child in root.findall(f"{{{ct_ns}}}Override"):
                    partname = child.attrib.get("PartName", "")
                    if partname == "/word/footnotes.xml":
                        overrides_to_remove.append(child)
                for child in overrides_to_remove:
                    root.remove(child)
                endnotes_exists = any(
                    child.attrib.get("PartName", "") == "/word/endnotes.xml"
                    for child in root.findall(f"{{{ct_ns}}}Override")
                )
                if not endnotes_exists:
                    endnote_override = ET.Element(
                        f"{{{ct_ns}}}Override",
                        PartName="/word/endnotes.xml",
                        ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"
                    )
                    root.append(endnote_override)
                updated_data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                updated_files[file] = updated_data

            elif file == "word/_rels/document.xml.rels":
                ns_rel = {"pr": "http://schemas.openxmlformats.org/package/2006/relationships"}
                tree = ET.ElementTree(ET.fromstring(data))
                root = tree.getroot()
                for rel in root.findall("pr:Relationship", ns_rel):
                    if rel.attrib.get("Type") == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes":
                        rel.attrib["Type"] = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes"
                        if rel.attrib.get("Target") == "footnotes.xml":
                            rel.attrib["Target"] = "endnotes.xml"
                updated_data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
                updated_files[file] = updated_data

            else:
                # For all other files, retain the original data.
                updated_files[file] = data

    # Write all updated (and unchanged) files into the output DOCX
    with zipfile.ZipFile(output_docx, 'w', zipfile.ZIP_DEFLATED) as zout:
        for file, data in updated_files.items():
            zout.writestr(file, data)


def write_to_log(doc_id, user):
    global global_logs
    current_date = datetime.now().strftime("%Y-%m-%d")
    output_path_file = Path(os.getcwd()) / 'output' / user / current_date / str(doc_id) / 'text'
    # dir_path = output_path_file.parent
    # output_path_file = Path(os.getcwd())/'output'/user
    # print(output_path_file)
    # output_dir = os.path.join('output', str(doc_id))
    os.makedirs(output_path_file, exist_ok=True)
    log_file_path = os.path.join(output_path_file, 'global_logs.txt')

    with open(log_file_path, 'a', encoding='utf-8') as log_file:
        log_file.write("\n".join(global_logs))

    global_logs = []
    


def process_doc_function10(payload: dict, doc: Document, doc_id, userNew):
    conn = get_db_connection()
    if conn is None:
        raise HTTPException(status_code=500, detail="Database connection error")
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM row_document WHERE row_doc_id = %s", (doc_id,))
    rows = cursor.fetchone()
    # print(rows)
    user_id=rows[5]
    cursor.execute("SELECT admin_name from admins where admin_id = %s",(user_id,))
    user = cursor.fetchone()
    conn.close()
    
    
    if not rows:
        raise HTTPException(status_code=404, detail="Document not found")
    username = user[0]
    
    current_date = datetime.now().strftime("%Y-%m-%d")
    output_path = os.path.join(os.getcwd(), 'output', username, current_date, str(doc_id), 'doc', "footnotes.docx")
    
    file_path = os.path.join(os.getcwd(), 'files', rows[1])
    footnotes = extract_footnotes(file_path)
    
    if footnotes:
        write_footnotes_to_docx(footnotes, output_path)
    
    endnote_path = str(os.path.join(os.getcwd(), 'output', username, current_date, str(doc_id), 'doc', "end_notes_document.docx"))
    convert_footnotes_to_endnotes(file_path, endnote_path)
    line_number = 1

        
    write_to_log(doc_id, username)