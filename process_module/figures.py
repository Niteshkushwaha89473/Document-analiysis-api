import re
from docx import Document
import os
import roman
from word2number import w2n
from pathlib import Path
from datetime import datetime

global_logs = []

def remove_dot_in_figure_number(runs):
   
    pattern = re.compile(r'(Figure\s+\d+\.\d+)\.')
    
    for run in runs:
        new_text = pattern.sub(r'\1', run.text)
        if new_text != run.text:
            run.text = new_text



def add_dot_to_figure_caption(runs):
    if not runs:
        return

    full_text = "".join(run.text for run in runs).strip()

    if not re.match(r'^Figure\s+\d+(?:\.\d+)*', full_text, re.IGNORECASE):
        return

    if full_text.endswith("."):
        return

    new_text = full_text + "."

    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""





def fix_figure_caption_format(runs):
    if not runs:
        return

    full_text = "".join(run.text for run in runs).strip()

    pattern = re.compile(r'^(?:FIGURE|Figure|fig\.|Fig\.|fig|figure)\s*(\d+(?:\.\d+)*)(?:\s+)?(.*)$')
    match = pattern.match(full_text)
    if not match:
        return

    number = match.group(1)
    remainder = match.group(2)

    # Construct the new parts.
    new_run1_text = "Figure"
    new_run2_text = " " + number + " "
    new_run3_text = remainder

    paragraph = runs[0]._parent

    for run in runs:
        run.text = ""

    run1 = paragraph.add_run(new_run1_text)
    run1.bold = True
    run1.font.name = "Calibri"

    run2 = paragraph.add_run(new_run2_text)
    run2.font.name = "Calibri"

    if new_run3_text:
        run3 = paragraph.add_run(new_run3_text)
        run3.font.name = "Calibri"
        # paragraph.add_run(new_run3_text)

        

def write_to_log(doc_id, user):
    global global_logs
    current_date = datetime.now().strftime("%Y-%m-%d")
    output_path_file = Path(os.getcwd()) / 'output' / user / current_date / str(doc_id) / 'text' 
    os.makedirs(output_path_file, exist_ok=True)
    log_file_path = os.path.join(output_path_file, 'global_logs.txt')

    with open(log_file_path, 'a', encoding='utf-8') as log_file:
        log_file.write("\n".join(global_logs))

    global_logs = []



def process_doc_function8(payload: dict, doc: Document, doc_id, user):
    for para in doc.paragraphs:
        fix_figure_caption_format(para.runs)
        remove_dot_in_figure_number(para.runs)
        add_dot_to_figure_caption(para.runs)
        
    write_to_log(doc_id, user)