import os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn

global_logs = []



# def process_quotes(doc):
#     """
#     Processes paragraphs in a Word document:
#     - If a paragraph is followed by one that starts with '–', process the first paragraph.
#     - Remove surrounding double quotes from the paragraph.
#     - Remove italics formatting from the paragraph.
#     """
#     paragraphs = doc.paragraphs
#     for i in range(len(paragraphs) - 1):
#         current_para = paragraphs[i]
#         next_para = paragraphs[i + 1]

#         if next_para.text.strip().startswith("–"):
#             for run in current_para.runs:
#                 if run.text.startswith('“') and run.text.endswith('”'):
#                     run.text = run.text[1:-1]
                
#                 if run.italic:
#                     run.italic = False

def process_quotes(doc):
    """
    Processes paragraphs in a Word document:
    - If a paragraph is followed by one that starts with '–', process the first paragraph.
    - Remove surrounding double quotes from the paragraph.
    - Remove italics formatting from the paragraph.
    """
    paragraphs = doc.paragraphs
    for i in range(len(paragraphs) - 1):
        current_para = paragraphs[i]
        next_para = paragraphs[i + 1]

        if next_para.text.strip().startswith("–"):
            for run in current_para.runs:
                # if run.text.startswith('“') and run.text.endswith('”'):
                if run.text.startswith('“'):
                    run.text = run.text[1:]
                if run.text.endswith('”'):
                    run.text = run.text[:-1]
                if run.italic:
                    run.italic = False



from docx.enum.text import WD_ALIGN_PARAGRAPH

def right_align_dash_paragraphs(doc):
    """
    Right-aligns paragraphs that start with an em dash ("–").
    """
    for para in doc.paragraphs:
        if para.text.strip().startswith("–"):
            # para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            para.alignment = 2



def write_to_log(doc_id, user):
    global global_logs
    current_date = datetime.now().strftime("%Y-%m-%d")
    output_path_file = Path(os.getcwd()) / 'output' / user / current_date / str(doc_id) / 'text' 
    # dir_path = output_path_file.parent

    # output_dir = os.path.join('output', str(doc_id))
    os.makedirs(output_path_file, exist_ok=True)
    log_file_path = os.path.join(output_path_file, 'global_logs.txt')

    with open(log_file_path, 'a', encoding='utf-8') as log_file:
        log_file.write("\n".join(global_logs))

    global_logs = []




def process_doc_function12(payload: dict, doc: Document, doc_id, user):
    process_quotes(doc)
    right_align_dash_paragraphs(doc)

    write_to_log(doc_id, user)