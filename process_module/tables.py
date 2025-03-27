import re
from docx import Document
import os
import roman
from word2number import w2n
from pathlib import Path
from datetime import datetime
from docx.enum.text import WD_ALIGN_PARAGRAPH

global_logs = []


def set_table_text_alignment(document):
    """
    For each cell in every table of the document:
      - If the cell's paragraph text is numeric, center-align the text.
        If the numeric value is a whole number, reformat it as a float with one decimal (e.g. '10' -> '10.0').
      - Otherwise, left-align the text.
    """
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    stripped_text = para.text.strip()
                    try:
                        # Attempt to convert the text to a float.
                        number = float(stripped_text)
                        # If conversion is successful, center-align the paragraph.
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # If the number is whole, format with one decimal place.
                        if number.is_integer():
                            new_text = f"{number:.1f}"
                        else:
                            new_text = stripped_text

                        # Update the paragraph's text. 
                        # We update the first run and clear out the remaining ones.
                        if para.runs:
                            para.runs[0].text = new_text
                            for run in para.runs[1:]:
                                run.text = ""
                        else:
                            para.add_run(new_text)
                    except ValueError:
                        # Not numeric; left-align the paragraph.
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    
                    
def format_dashes_in_tables(document):
    """
    In each table cell, if a paragraph contains only a dash ('-')
    indicating no data, replace it with an em dash ('—') and center it.
    """
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    # Check if the paragraph's text (stripped of whitespace) is exactly a dash
                    if para.text.strip() == "-":
                        # Replace normal dash with em dash in each run
                        for run in para.runs:
                            run.text = run.text.replace("-", "—")
                        # Center the paragraph
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER




previous_end = None

def process_table_spans(runs):
    if not runs:
        return

    # Extract text from runs
    text = ''.join(run.text for run in runs)
    
    # Regular expression to find number spans like 0–100, 100–200, etc.
    pattern = re.compile(r'\b(\d+)\s*[–-]\s*(\d+)\b')
    matches = list(pattern.finditer(text))
    
    if not matches:
        return

    global previous_end
    new_text = text

    for match in matches:
        start, end = int(match.group(1)), int(match.group(2))
        if previous_end is not None:
            if start <= previous_end:
                start = previous_end + 1
                replacement = f"{start}–{end}"
                new_text = new_text.replace(match.group(0), replacement, 1)
            else:
                replacement = match.group(0)
        else:
            replacement = match.group(0)

        previous_end = end
    
    # Update runs text
    offset = 0
    for run in runs:
        length = len(run.text)
        run.text = new_text[offset:offset + length]
        offset += length


def process_document_tables_ranges(doc):
    global previous_end
    for table in doc.tables:
        for row in table.rows:
            # Process only the first column
            first_cell = row.cells[0]
            for para in first_cell.paragraphs:
                process_table_spans(para.runs)
                
    previous_end = None
    
    
    
    
def remove_dot_in_table_number(runs):
    """
    Removes the trailing dot from Table numbers (e.g., "Table 1.1." -> "Table 1.1").
    """
    pattern = re.compile(r'(Table\s+\d+\.\d+)\.')
    
    for run in runs:
        new_text = pattern.sub(r'\1', run.text)
        if new_text != run.text:
            run.text = new_text


def add_dot_to_table_caption(runs):
    """
    Ensures that a Table caption ends with a period.
    """
    if not runs:
        return

    full_text = "".join(run.text for run in runs).strip()

    if not re.match(r'^Table\s+\d+(?:\.\d+)*', full_text, re.IGNORECASE):
        return

    if full_text.endswith("."):
        return

    new_text = full_text + "."

    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""


def fix_table_caption_format(runs):
    """
    Standardizes the format of Table captions to "Table X.Y Description" with proper styling.
    """
    if not runs:
        return

    full_text = "".join(run.text for run in runs).strip()

    pattern = re.compile(r'^(?:TABLE|Table|tab\.|Tab\.|tab|table)\s*(\d+(?:\.\d+)*)(?:\s+)?(.*)$')
    match = pattern.match(full_text)
    if not match:
        return

    number = match.group(1)
    remainder = match.group(2)

    # Construct the new parts.
    new_run1_text = "Table"
    new_run2_text = " " + number + " "
    new_run3_text = remainder

    paragraph = runs[0]._parent

    for run in runs:
        run.text = ""

    run1 = paragraph.add_run(new_run1_text)
    run1.bold = True
    run1.font.name = "Calibri"

    run2 = paragraph.add_run(new_run2_text)
    run2.font.name = "Calibri"

    if new_run3_text:
        run3 = paragraph.add_run(new_run3_text)
        run3.font.name = "Calibri"




def write_to_log(doc_id, user):
    global global_logs
    current_date = datetime.now().strftime("%Y-%m-%d")
    output_path_file = Path(os.getcwd()) / 'output' / user / current_date / str(doc_id) / 'text' 
    # dir_path = output_path_file.parent

    # output_dir = os.path.join('output', str(doc_id))
    os.makedirs(output_path_file, exist_ok=True)
    log_file_path = os.path.join(output_path_file, 'global_logs.txt')

    with open(log_file_path, 'a', encoding='utf-8') as log_file:
        log_file.write("\n".join(global_logs))

    global_logs = []




def process_doc_function9(payload: dict, doc: Document, doc_id, user):
    set_table_text_alignment(doc)
    format_dashes_in_tables(doc)
    process_document_tables_ranges(doc)
    for para in doc.paragraphs:
        fix_table_caption_format(para.runs)
        add_dot_to_table_caption(para.runs)
        remove_dot_in_table_number(para.runs)
        
    write_to_log(doc_id, user)