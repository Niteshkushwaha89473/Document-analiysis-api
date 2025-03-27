import re
from docx import Document
import os
from db_config import get_db_connection
from datetime import datetime
from docx.text.run import Run
from pathlib import Path

# Global logs to keep track of changes
global_logs = []
# seen_symbols = set()


# A map of numbers to century strings
century_map = {
    1: "first",
    2: "second",
    3: "third",
    4: "fourth",
    5: "fifth",
    6: "sixth",
    7: "seventh",
    8: "eighth",
    9: "ninth",
    10: "tenth",
    11: "eleventh",
    12: "twelfth",
    13: "thirteenth",
    14: "fourteenth",
    15: "fifteenth",
    16: "sixteenth",
    17: "seventeenth",
    18: "eighteenth",
    19: "nineteenth",
    20: "twentieth",
    21: "twenty-first",
    22: "twenty-second",
    23: "twenty-third",
    24: "twenty-fourth",
    25: "twenty-fifth",
}


def fetch_abbreviation_mappings():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT original_word, abbreviated_form FROM abbreviation_mapping")
    mappings = cursor.fetchall()
    conn.close()
    return {row[0]: row[1] for row in mappings}

# change long form to short form abbrevaited form
# def apply_abbreviation_mapping(runs, abbreviation_dict, line_number):
#     global global_logs
#     for run in runs:
#         words = run.text.split()
#         updated_words = []
        
#         for word in words:
#             updated_word = abbreviation_dict.get(word, word)
#             if word != updated_word:
#                 global_logs.append(f"[apply_abbreviation_mapping] Line {line_number}: '{word}' -> '{updated_word}'")
#             updated_words.append(updated_word)
        
#         # Join the updated words and reassign it to the run's text
#         run.text = ' '.join(updated_words)


def apply_abbreviation_mapping(runs, abbreviation_dict, line_number):
    global global_logs

    for run in runs:
        original_text = run.text

        def replace_match(match):
            word = match.group(0)  # Extract matched word
            stripped_word = word.rstrip('.,')  # Remove trailing punctuation
            
            updated_word = abbreviation_dict.get(stripped_word, stripped_word)  # Find replacement
            
            if stripped_word != updated_word:
                global_logs.append(f"[apply_abbreviation_mapping] Line {line_number}: '{stripped_word}' -> '{updated_word}'")

            # Preserve punctuation
            return updated_word + word[len(stripped_word):]

        # Use regex to match full words with optional punctuation
        pattern = r'\b[\w]+\b'
        updated_text = re.sub(pattern, replace_match, original_text)

        # Assign the corrected text back
        run.text = updated_text





# Converts century notation like '21st' to 'the twenty-first century'
# def convert_century_in_runs(runs, line_number_offset):
#     """
#     Converts century notation like '21st' to 'the twenty-first century'
#     and logs the changes with line numbers.

#     :param runs: List of Run objects in the paragraph.
#     :param line_number_offset: The starting line number for this paragraph.
#     :return: None (modifies runs in place).
#     """
#     global global_logs  # Global log to record changes

#     for run in runs:
#         words = run.text.split()  # Split the run's text into words
#         updated_words = []

#         for word in words:
#             match = re.match(r"(\d+)(st|nd|rd|th)$", word)  # Match century notation
#             if match:
#                 num = int(match.group(1))
#                 if num in century_map:
#                     # Original and converted word
#                     original_word = match.group(0)
#                     converted_word = f"the {century_map[num]}"
                    
#                     # Log the change with the actual line number
#                     global_logs.append(
#                         f"[convert century] Line {line_number_offset}: {original_word} -> {converted_word}"
#                     )
                    
#                     # Replace the word in the run
#                     word = converted_word
            
#             updated_words.append(word)
        
#         # Rebuild the run's text with updated words
#         run.text = ' '.join(updated_words)


# Converts century notation like '21st' to 'the twenty-first century'
def convert_century_in_runs(runs, line_number_offset):
    """
    Converts century notation like '21st' to 'the twenty-first century'
    and logs the changes with line numbers.

    :param runs: List of Run objects in the paragraph.
    :param line_number_offset: The starting line number for this paragraph.
    :return: None (modifies runs in place).
    """
    for run in runs:
        new_text = run.text  # Store the original text
        # Match century notation within the text
        matches = re.finditer(r"(\d+)(st|nd|rd|th)\b", run.text)
        for match in matches:
            num = int(match.group(1))
            if num in century_map:
                original_word = match.group(0)
                converted_word = f"{century_map[num]}"
                
                # Log the change
                global_logs.append(
                    f"[convert century] Line {line_number_offset}: {original_word} -> {converted_word}"
                )
                
                # Replace occurrences in the text
                new_text = new_text.replace(original_word, converted_word)
        
        # Update run text
        run.text = new_text



# change italics of latin word to roman
# def set_latinisms_to_roman_in_runs(runs, line_number, latinisms=None):
#     """
#     Converts specific Latinisms from italic to roman text in a paragraph.
#     Logs changes to the global_log, including line number and original italicized Latinism.

#     :param runs: List of Run objects in the paragraph.
#     :param line_number: The line number for logging purposes.
#     :param latinisms: List of Latinisms to convert (defaults to common Latinisms).
#     :return: None (modifies runs in place).
#     """
#     if latinisms is None:
#         latinisms = [
#             "i.e.", "e.g.", "via", "vice versa", "etc.", "a posteriori", 
#             "a priori", "et al.", "cf.", "c."
#         ]
#     global global_logs

#     for run in runs:
#         for lat in latinisms:
#             if lat in run.text:
#                 # Log the change
#                 global_logs.append(
#                     f"[set_latinisms_to_roman] Line {line_number}: '{lat}' -> '{lat}'"
#                 )
                
#                 # Remove italic formatting for the Latinism
#                 run.text = run.text.replace(lat, lat)
#                 run.font.italic = False  # Set the font to non-italic



# change italics of latin word to roman
def set_latinisms_to_roman_in_runs(runs, line_number, latinisms=None):
    """
    Converts specific Latinisms from italic to roman text in a paragraph.
    Logs changes to the global_logs, including line number and original italicized Latinism.

    :param runs: List of Run objects in the paragraph.
    :param line_number: The line number for logging purposes.
    :param latinisms: List of Latinisms to convert (defaults to common Latinisms).
    :return: None (modifies runs in place).
    """
    if latinisms is None:
        latinisms = [
            "i.e.", "e.g.", "via", "vice versa", "etc.", "a posteriori",
            "a priori", "et al.", "cf.", "c."
        ]
    # Sort Latinisms by descending length to handle longer matches first
    latinisms = sorted(latinisms, key=lambda x: len(x), reverse=True)
    
    global global_logs
    if not runs:
        return
    
    parent = runs[0]._parent  # Get the parent paragraph
    
    # Process runs in reverse to avoid issues with modifying the list while iterating
    for i in range(len(runs) - 1, -1, -1):
        run = runs[i]
        original_text = run.text
        if not original_text:
            continue
        
        # Collect all occurrences of any Latinism in the run's text
        split_points = []
        for lat in latinisms:
            start = 0
            while start <= len(original_text):
                pos = original_text.find(lat, start)
                if pos == -1:
                    break
                # Check if the found occurrence is a whole word (to avoid partial matches)
                # Ensure it's surrounded by word boundaries or string edges
                if (pos == 0 or not original_text[pos-1].isalnum()) and \
                   (pos + len(lat) == len(original_text) or not original_text[pos + len(lat)].isalnum()):
                    split_points.append((pos, pos + len(lat), lat))
                start = pos + 1  # Move past this occurrence
        
        if not split_points:
            continue  # No Latinism found in this run
        
        # Sort split points by their start position
        split_points.sort(key=lambda x: x[0])
        
        # Split the text into parts
        parts = []
        last_end = 0
        for start, end, lat in split_points:
            if start > last_end:
                parts.append(original_text[last_end:start])
            parts.append((original_text[start:end], lat))
            last_end = end
        if last_end < len(original_text):
            parts.append(original_text[last_end:])
        
        # Create new runs for each part
        new_elements = []
        for part in parts:
            if isinstance(part, tuple):
                part_text, lat = part
                # Create a new run with the same style but italic False
                new_run = parent.add_run()
                new_run.text = part_text
                # Copy font properties from original run
                original_font = run.font
                new_font = new_run.font
                new_font.italic = False
                new_font.bold = original_font.bold
                new_font.underline = original_font.underline
                if original_font.size:
                    new_font.size = original_font.size
                if original_font.name:
                    new_font.name = original_font.name
                # Log the change (avoid duplicates)
                log_entry = f"[set_latinisms_to_roman] Line {line_number}: '{lat}' -> '{lat}'"
                if log_entry not in global_logs:
                    global_logs.append(log_entry)
            else:
                # Create a new run with the same style as original
                new_run = parent.add_run()
                new_run.text = part
                original_font = run.font
                new_font = new_run.font
                new_font.italic = original_font.italic
                new_font.bold = original_font.bold
                new_font.underline = original_font.underline
                if original_font.size:
                    new_font.size = original_font.size
                if original_font.name:
                    new_font.name = original_font.name
            new_elements.append(new_run._element)
        
        # Replace the original run with the new elements
        run_element = run._element
        parent_element = parent._element
        index = parent_element.index(run_element)
        parent_element.remove(run_element)
        # Insert new elements in reverse to maintain order
        for elem in reversed(new_elements):
            parent_element.insert(index, elem)
            


# make symbols for copyright only once
def process_symbols_mark_in_runs(runs, line_number, seen_symbols, symbols=None):
    """
    Ensures symbols like ®, ™, etc., appear only the first time in the text.
    Updates the global_log with changes, including line number, original text, and updated text.

    :param runs: List of Run objects in the paragraph.
    :param line_number: The line number for logging purposes.
    :param symbols: List of symbols to process (defaults to ["®", "™", "©", "℗", "℠"]).
    :return: None (modifies runs in place).
    """
    if symbols is None:
        symbols = ["®", "™", "©", "℗", "℠"]

    global global_logs
    
    updated_runs = []

    for run in runs:
        original_text = run.text
        new_text = original_text

        for symbol in symbols:
            if symbol in new_text:
                if symbol in seen_symbols:
                    # Remove all further occurrences
                    new_text = new_text.replace(symbol, "")
                else:
                    seen_symbols.add(symbol)  # Mark as first occurrence

        if new_text != original_text:
            run.text = new_text  # Apply changes
            updated_runs.append(symbol)

    # Log changes if any symbols were removed
    if updated_runs:
        global_logs.append(
            f"[process_symbols_in_doc] Line {line_number}: Removed duplicate symbols {updated_runs}"
        )


# change italics of see to roman
def apply_remove_italics_see_rule_in_runs(runs):
    """
    Replaces '*see*' with 'see' in the text, removing italics for the word 'see'.
    Modifies the runs in place.

    :param runs: List of Run objects in the paragraph.
    :return: None (modifies runs in place).
    """
    for run in runs:
        if '*see*' in run.text:
            # Replace '*see*' with 'see' and remove italics for this run
            run.text = run.text.replace('*see*', 'see')
            run.font.italic = False  # Ensure 'see' is not italicized





# change number to no. if followed by number
def set_number_to_no_in_runs(runs, line_number):
    """
    Replaces 'Number X' or 'number X' with 'No. X' or 'no. X' and logs changes.
    Modifies the runs in place.

    :param runs: List of Run objects in the paragraph.
    :param line_number: Line number for logging.
    :return: None (modifies runs in place).
    """
    global global_logs

    def replace_number(match):
        word = match.group(1)
        num = match.group(2)
        updated_text = f"No. {num}" if word.istitle() else f"no. {num}"
        global_logs.append(f"[set_number_to_no] Line {line_number}: '{match.group(0)}' -> '{updated_text}'")
        return updated_text

    pattern = r'\b(Number|number)\s(\d+)\b'

    for run in runs:
        if re.search(pattern, run.text):
            # Replace 'Number X' or 'number X' with 'No. X' or 'no. X'
            run.text = re.sub(pattern, replace_number, run.text)



# change long titile before name to short form
def format_titles_us_english_with_logging_in_runs(runs, line_number):
    """
    Replaces long titles (e.g., 'Doctor', 'Mister') with their short forms (e.g., 'Dr.', 'Mr.').
    Logs changes to the global_log.
    :param runs: List of Run objects in the paragraph.
    :return: None (modifies runs in place).
    """
    global global_logs
    titles = {
        "doctor": "Dr.",
        "mister": "Mr.",
        "misses": "Mrs.",
        "miss": "Miss.",
        "ms": "Ms.",
        "professor": "Professor",
        "sir": "Sir",
        "madam": "Madam",
        "saint": "St",
        "Doctor": "Dr.",
        "Mister": "Mr.",
        "Misses": "Mrs.",
        "Miss": "Miss.",
        "Ms": "Ms.",
        "Professor": "Professor",
        "Sir": "Sir",
        "Madam": "Madam",
        "Saint": "St",
    }

    for run in runs:
        for title, replacement in titles.items():
            if title in run.text:
                # Replace the title with its short form
                run.text = re.sub(rf"\b{title}\b", replacement, run.text, flags=re.IGNORECASE)
                # Log the change
                global_logs.append(f"[shorten title] Line {line_number}: {title} -> {replacement}")



# a.m. not AM or am.
# p.m not PM or pm.
# def enforce_am_pm_in_runs(runs, line_num):
#     """
#     Ensures consistent formatting for 'am' and 'pm' in the paragraph and logs changes.
#     Modifies the runs in place.

#     :param runs: List of Run objects in the paragraph.
#     :param line_num: The line number in the document for logging.
#     :return: None (modifies runs in place).
#     """
#     global global_logs

#     # Combine all run texts into one string
#     full_text = "".join(run.text for run in runs)

#     # Define pattern to match different variations of am/pm
#     pattern = r'\b(1[0-2]|0?[1-9])\s?(am|a\.m|pm|p\.m)\b'
#     # pattern = r'\b(1[0-2]|0?[1-9])\s?(AM|PM|A\.M\.|P\.M\.|Am|Pm|A\.m\.|P\.m\.|a\.m\.|p\.m\.|am|pm|AM\,)\.?\b'


#     def replace_am_pm(match):
#         time_value, period = match.groups()
#         corrected_period = "a.m." if "a" in period.lower() else "p.m."
#         corrected_text = f"{time_value} {corrected_period}"
#         global_logs.append(f"[am pm change] Line {line_num}: '{match.group(0)}' -> '{corrected_text}'")
#         return corrected_text

#     # Apply corrections
#     corrected_text = re.sub(pattern, replace_am_pm, full_text, flags=re.IGNORECASE)

#     # Reassign corrected text back to runs while preserving formatting
#     remaining_text = corrected_text
#     for run in runs:
#         run_length = len(run.text)
#         run.text = remaining_text[:run_length]
#         remaining_text = remaining_text[run_length:]


def enforce_am_pm_in_runs(runs, line_num):
    """
    Ensures consistent formatting for 'am' and 'pm' in the paragraph and logs changes.
    Modifies the runs in place.

    :param runs: List of Run objects in the paragraph.
    :param line_num: The line number in the document for logging.
    :return: None (modifies runs in place).
    """
    global global_logs  # Use a global log to record changes

    # Define pattern to match different variations of am/pm
    pattern = r'\b(1[0-2]|0?[1-9])\s?(am|a\.m\.?|pm|p\.m\.?)\b'

    def replace_am_pm(match):
        time_value, period = match.groups()
        corrected_period = "a.m." if "a" in period.lower() else "p.m."
        corrected_text = f"{time_value} {corrected_period}"
        global_logs.append(f"[am pm change] Line {line_num}: '{match.group(0)}' -> '{corrected_text}'")
        return corrected_text

    for run in runs:
        original_text = run.text
        corrected_text = re.sub(pattern, replace_am_pm, original_text, flags=re.IGNORECASE)

        if corrected_text != original_text:
            run.text = corrected_text  # Modify run in place



# apples, pears, and bananas
# apples, pears, or bananas
# def enforce_serial_comma_in_runs(runs):
#     """
#     Ensures the use of the serial comma (Oxford comma) in lists.
#     Modifies the runs in place.

#     :param runs: List of Run objects in the paragraph.
#     :return: None (modifies runs in place).
#     """
#     for run in runs:
#         # Add a comma before "and" or "or" in lists
#         run.text = re.sub(
#             r'([^,]+), ([^,]+) (or) ([^,]+)',
#             r'\1, \2, \3 \4',
#             run.text
#         )
#         # Explicitly handle cases where "and" does not get the serial comma
#         run.text = re.sub(
#             r'([^,]+), ([^,]+) (and) ([^,]+)',
#             r'\1, \2, \3 \4',
#             run.text
#         )


def enforce_serial_comma_in_runs(runs):
    """
    Ensures the use of the serial comma (Oxford comma) in lists.
    Modifies the runs in place.

    :param runs: List of Run objects in the paragraph.
    :return: None (modifies runs in place).
    """
    for run in runs:
        # Add a comma before "and" or "or" in lists only when necessary
        run.text = re.sub(
            r'([^,]+), ([^,]+) (and|or) ([^,]+)',
            r'\1, \2, \3 \4',
            run.text
        )
        # Fix the case for when "and" and "or" are at the end
        # run.text = re.sub(
        #     r'([^,]+) (and|or) ([^,]+)',
        #     r'\1, \2 \3',
        #     run.text
        # )




# Replace all occurrences of the § symbol with 'Section'
def rename_section_in_runs(runs):
    """
    Replaces all occurrences of the § symbol with 'Section'.
    Modifies the runs in place.

    :param runs: List of Run objects in the paragraph.
    :return: None (modifies runs in place).
    """
    for run in runs:
        run.text = re.sub(r'§', 'Section', run.text)


# There is one problem here for project, & document it is not changing and for project & document it is changing
# chnage and to & if in between two word starting with capitals
def replace_ampersand_in_runs(runs, line_number):
    """
    Replaces '&' with 'and' unless both sides are uppercase (e.g., 'R&D').
    Modifies the runs in place and logs changes.

    :param runs: List of Run objects in the paragraph.
    :return: None (modifies runs in place).
    """
    global global_logs

    def replacement(match):
        left, right = match.group(1), match.group(2)
        original = match.group(0)
        if left[0].isupper() and right[0].isupper():
            return original  # Preserve '&' if both sides are uppercase (e.g., 'R&D')
        modified = left + ' and ' + right
        global_logs.append(
            f"[replace_ampersand] Line {line_number}: '{original}' -> '{modified}'"
        )
        return modified

    for run in runs:
        run.text = re.sub(r'(\w+)\s*&\s*(\w+)', replacement, run.text)



# changes word like James' to James's
def correct_possessive_names_in_runs(runs, line_number):
    """
    Corrects possessive names (e.g., "James'" to "James's").
    Modifies the runs in place and logs changes.

    :param runs: List of Run objects in the paragraph.
    :param line_number: The line number for logging.
    :return: None (modifies runs in place).
    """
    global global_logs

    for run in runs:
        # Correct singular possessive (e.g., "James'" to "James's")
        run.text = re.sub(r"\b([A-Za-z]+s)\b(?<!\bs')'", r"\1's", run.text)
        
        # Correct plural possessive (e.g., "students'" remains "students'")
        run.text = re.sub(r"\b([A-Za-z]+s)'\b", r"\1'", run.text)




def remove_and_in_runs(runs, line_number):
    """
    Replaces 'and' between two capitalized words with an ampersand (&).
    Modifies the runs in place and logs changes.

    :param runs: List of Run objects in the paragraph.
    :param line_number: The line number for logging.
    :return: None (modifies runs in place).
    """
    global global_logs

    pattern = r'([A-Z][a-z]+)\s+and\s+([A-Z][a-z]+)'

    def process_and_replacement(match):
        original = match.group(0)
        modified = f"{match.group(1)} & {match.group(2)}"
        if original != modified:
            global_logs.append(
                f"[remove_and] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified

    for run in runs:
        run.text = re.sub(pattern, process_and_replacement, run.text)
        
        
def remove_quotation_in_runs(runs, line_number):
    """
    Removes single quotation marks (') following capitalized words.
    Modifies the runs in place and logs changes.

    :param runs: List of Run objects in the paragraph.
    :param line_number: The line number for logging.
    :return: None (modifies runs in place).
    """
    global global_logs

    pattern = r"([A-Z]+)[’']"

    def process_quotation_removal(match):
        original = match.group(0)
        modified = f"{match.group(1)}"
        if original != modified:
            global_logs.append(
                f"[remove_quotation] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified

    for run in runs:
        run.text = re.sub(pattern, process_quotation_removal, run.text)
        
        




def correct_acronyms_in_runs(runs, line_number):
    """
    Removes periods from acronyms (e.g., 'U.S.A.' becomes 'USA') using regex substitution.
    Modifies the runs in place and logs changes without altering original spacing.

    :param runs: List of Run objects in the paragraph.
    :param line_number: The line number for logging.
    :return: None (modifies runs in place).
    """
    global global_logs

    for run in runs:
        original_text = run.text

        # Function to replace acronyms and log changes
        def replace_acronym(match, is_upper):
            original = match.group()
            # Check if the acronym matches the correct pattern (either upper or lower case)
            if is_upper:
                if not re.fullmatch(r"([A-Z]\.){2,}[A-Z]\.?", original):
                    return original
            else:
                if not re.fullmatch(r"([a-z]\.){2,}[a-z]\.?", original):
                    return original
            replaced = original.replace(".", "")
            global_logs.append(
                f"[correct_acronyms] Line {line_number}: '{original}' -> '{replaced}'"
            )
            return replaced

        # Process uppercase acronyms as standalone words
        processed_text = re.sub(
            r"\b([A-Z]\.){2,}[A-Z]\.?\b",
            lambda m: replace_acronym(m, True),
            original_text
        )

        # Process lowercase acronyms as standalone words
        processed_text = re.sub(
            r"\b([a-z]\.){2,}[a-z]\.?\b",
            lambda m: replace_acronym(m, False),
            processed_text
        )

        if processed_text != original_text:
            run.text = processed_text




# changes eg to e.g.

def enforce_eg_rule_with_logging_in_runs(runs, line_number):
    global global_logs

    for run in runs:
        original_text = run.text

        # Step 1: Replace "eg," (with a comma) directly with "e.g."
        new_text = re.sub(r'\beg,\b', 'e.g.', original_text, flags=re.IGNORECASE)  

        # Step 2: Replace standalone "eg" (without a comma) with "e.g."
        new_text = re.sub(r'\beg\b', 'e.g.', new_text, flags=re.IGNORECASE)

        # Step 3: Fix extra periods like "e.g.." or "e.g..."
        new_text = re.sub(r'e\.g\.\.+', 'e.g.', new_text)

        # Step 4: Remove any trailing comma after "e.g." (e.g., "e.g.," -> "e.g.")
        new_text = re.sub(r'e\.g\.,', 'e.g.', new_text)

        # Log changes if any updates were made
        if new_text != original_text:
            global_logs.append(
                f"[e.g. correction] Line {line_number}: '{original_text.strip()}' -> '{new_text.strip()}'"
            )
        
        # Update the run's text
        run.text = new_text

        
# def enforce_eg_rule_with_logging_in_runs(runs, line_number):
#     """
#     Ensures consistent formatting for 'e.g.' in the paragraph and logs changes.
#     It normalizes different variants (like "eg", "eg," or "e.g,") into "e.g.",
#     fixes extra periods (e.g. "e.g..", "e.g...") and finally removes a trailing
#     comma if present (converting "e.g.," into "e.g.").
    
#     Modifies the runs in place.

#     :param runs: List of Run objects in the paragraph.
#     :param line_number: The line number for logging.
#     :return: None (modifies runs in place).
#     """
#     global global_logs

#     for run in runs:
#         original_text = run.text
#         # Step 1: Normalize "eg" or incorrect "e.g" variants to "e.g."
#         new_text = re.sub(r'\beg,\b', 'e.g.', original_text, flags=re.IGNORECASE)  # Handle "eg,"
#         new_text = re.sub(r'\beg\b', 'e.g.', new_text, flags=re.IGNORECASE)
#         new_text = re.sub(r'e\.g,', 'e.g.', new_text)  # Handle incorrect "e.g,"

#         # Step 2: Fix extra periods like "e.g.." or "e.g..."
#         new_text = re.sub(r'e\.g\.\.+', 'e.g.', new_text)

#         # Step 3: Remove a trailing comma after "e.g." (i.e. change "e.g.," to "e.g.")
#         new_text = re.sub(r'e\.g\.,', 'e.g.', new_text)

#         # Log changes if any
#         if new_text != original_text:
#             global_logs.append(
#                 f"[e.g. correction] Line {line_number}: '{original_text.strip()}' -> '{new_text.strip()}'"
#             )
        
#         # Update the run's text
#         run.text = new_text


def enforce_ie_rule_with_logging_in_runs(runs, line_number):
    global global_logs

    for run in runs:
        original_text = run.text

        # Step 1: Replace "ie," (with a comma) directly with "i.e."
        new_text = re.sub(r'\bie,\b', 'i.e.', original_text, flags=re.IGNORECASE)  

        # Step 2: Replace standalone "ie" (without a comma) with "i.e."
        new_text = re.sub(r'\bie\b', 'i.e.', new_text, flags=re.IGNORECASE)

        # Step 3: Fix extra periods like "i.e.." or "i.e..."
        new_text = re.sub(r'i\.e\.\.+', 'i.e.', new_text)

        # Step 4: Remove any trailing comma after "i.e." (i.e., "i.e.," -> "i.e.")
        new_text = re.sub(r'i\.e\.,', 'i.e.', new_text)

        # Log changes if any updates were made
        if new_text != original_text:
            global_logs.append(
                f"[i.e. correction] Line {line_number}: '{original_text.strip()}' -> '{new_text.strip()}'"
            )
        
        # Update the run's text
        run.text = new_text



# def enforce_ie_rule_with_logging_in_runs(runs, line_number):
#     """
#     Ensures consistent formatting for 'i.e.' in the paragraph and logs changes.
#     Modifies the runs in place.

#     :param runs: List of Run objects in the paragraph.
#     :param line_number: The line number for logging.
#     :return: None (modifies runs in place).
#     """
#     global global_logs

#     for run in runs:
#         original_text = run.text

#         # Step 1: Match "ie" or "i.e." with optional surrounding spaces and punctuation
#         new_text = re.sub(r'\bie\b', 'i.e.', run.text, flags=re.IGNORECASE)  # Handle standalone "ie"
#         new_text = re.sub(r'\bie,\b', 'i.e.', new_text, flags=re.IGNORECASE)  # Handle "ie,"

#         # Step 2: Fix extra periods like `i.e..` or `i.e...,` and ensure proper punctuation
#         new_text = re.sub(r'\.([.,])', r'\1', new_text)  # Removes an extra period before a comma or period
#         new_text = re.sub(r'\.\.+', '.', new_text)  # Ensures only one period after i.e.

#         # Step 3: Remove comma if i.e... is followed by it (i.e..., -> i.e.)
#         new_text = re.sub(r'i\.e\.,', 'i.e.', new_text)
        
#         # Step 4: Change i.e, to i.e.
#         new_text = re.sub(r'i\.e,', 'i.e.', new_text)

#         # Log changes if the text is updated
#         if new_text != original_text:
#             global_logs.append(
#                 f"[i.e. correction] Line {line_number}: '{original_text.strip()}' -> '{new_text.strip()}'"
#             )
        
#         # Update the run's text
#         run.text = new_text

        
        

def standardize_etc_in_runs(runs, line_number):
    global global_logs
    pattern = r'\b(e\.?tc|e\.t\.c|e\.t\.c\.|et\.?\s?c|et\s?c|etc\.?|etc|et cetera|etcetera|Etc\.?|Etc|‘and etc\.’|et\.?\s?cetera|etc\.?,?|etc\.?\.?|etc\,?\.?)\b'

    # Get full paragraph text
    full_text = "".join(run.text for run in runs)

    # Define replacement function
    def replace_etc(match):
        original = match.group(0)
        modified = "etc."
        if original != modified:
            global_logs.append(f"[etc. correction] Line {line_number}: '{original}' -> '{modified}'")
        return modified

    # Replace variations of "etc."
    updated_text = re.sub(pattern, replace_etc, full_text, flags=re.IGNORECASE)

    # Explicit fixes for "etc..", "etc.,"
    updated_text = re.sub(r'etc\.\.+', 'etc.', updated_text)
    updated_text = re.sub(r'etc\.,', 'etc.', updated_text)

    # Update runs while keeping formatting
    current_index = 0
    for i, run in enumerate(runs):
        run.text = updated_text[current_index : current_index + len(run.text)]
        current_index += len(run.text)


def insert_thin_space_between_number_and_unit_in_runs(runs, line_number):
    global global_logs
    thin_space = '\u2009'

    # Get full paragraph text
    full_text = "".join(run.text for run in runs)
    
    pattern = r"(\d+)(\s*)([a-zA-Z]+)(?!\s?°)"

    # Define replacement function
    def replace_number_unit(match):
        number = match.group(1)
        whitespace = match.group(2)
        unit = match.group(3)
        original_word = number + whitespace + unit
        updated_word = number + thin_space + unit

        if original_word != updated_word:
            global_logs.append(f"[insert_thin_space_between_number_and_unit] Line {line_number}: '{original_word}' -> '{updated_word}'")

        return updated_word

    # Apply regex replacement to full text
    updated_text = re.sub(pattern, replace_number_unit, full_text)

    # Update runs while keeping formatting
    current_index = 0
    for i, run in enumerate(runs):
        run.text = updated_text[current_index : current_index + len(run.text)]
        current_index += len(run.text)


import re
import spacy

nlp = spacy.load("en_core_web_sm")

def process_paragraph(runs, line_number):
    """
    Processes a paragraph to:
    1. Add a comma before 'e.g.' if there is no verb between 'e.g.' and the end of the sentence.
    2. Add a colon before 'i.e.' wherever it appears.
    
    Args:
        runs (list): List of Word paragraph runs to preserve formatting.
        line_number (int): Line number for logging.
    """
    global global_logs
    
    # Extract full paragraph text from runs
    full_text = "".join(run.text for run in runs)

    # Use NLP to analyze sentence structure
    doc = nlp(full_text)
    updated_sentences = []

    for sentence in doc.sents:
        sentence_text = sentence.text

        # Step 1: Handle 'e.g.'
        if "e.g." in sentence_text:
            eg_start_idx = sentence_text.find("e.g.")
            after_eg_text = sentence_text[eg_start_idx:]
            
            # Check if there is a verb after 'e.g.'
            after_eg_doc = nlp(after_eg_text)
            has_verb = any(token.pos_ == "VERB" for token in after_eg_doc)

            if not has_verb:
                new_sentence_text = re.sub(r"(?<!,)\s+e\.g\.", ", e.g.", sentence_text)
                if new_sentence_text != sentence_text:
                    global_logs.append(f"[process_paragraph] Line {line_number}: Added comma before 'e.g.'")
                    sentence_text = new_sentence_text

        # Step 2: Handle 'i.e.'
        if "i.e." in sentence_text:
            new_sentence_text = re.sub(r"(?<!;)\s+i\.e\.", ": i.e.", sentence_text)
            if new_sentence_text != sentence_text:
                global_logs.append(f"[process_paragraph] Line {line_number}: Replaced space before 'i.e.' with ':'")
                sentence_text = new_sentence_text

        updated_sentences.append(sentence_text)

    updated_text = " ".join(updated_sentences)

    # Update runs while keeping formatting intact
    current_index = 0
    for i, run in enumerate(runs):
        run.text = updated_text[current_index : current_index + len(run.text)]
        current_index += len(run.text)


def apply_quotation_punctuation_rule(text: str):
    """
    Adjusts the placement of punctuation marks within single quotes.
    Args:
        text (str): Input text to process.
    Returns:
        str: Updated text.
    """
    global global_logs
    pattern = r"‘(.*?)’([!?])"
    def process_quotation_punctuation(match):
        original = match.group(0)
        modified = f"‘{match.group(1)}{match.group(2)}’"
        if original != modified:
            line_number = text[:match.start()].count('\n') + 1
            global_logs.append(
                f"[apply_quotation_punctuation_rule] Line {line_number}: '{original}' -> '{modified}'"
            )
        return modified
    updated_text = re.sub(pattern, process_quotation_punctuation, text)
    return updated_text



# Dictionary to convert word numbers to integer values
word_to_num = {
    'zero': 0, 'one': 1, 'two': 2, 'three': 3, 'four': 4, 'five': 5,
    'six': 6, 'seven': 7, 'eight': 8, 'nine': 9, 'ten': 10, 'eleven': 11,
    'twelve': 12, 'thirteen': 13, 'fourteen': 14, 'fifteen': 15, 'sixteen': 16,
    'seventeen': 17, 'eighteen': 18, 'nineteen': 19, 'twenty': 20, 'twenty-one': 21,
    'twenty-two': 22, 'twenty-three': 23, 'twenty-four': 24, 'twenty-five': 25,
    'twenty-six': 26, 'twenty-seven': 27, 'twenty-eight': 28, 'twenty-nine': 29,
    'thirty': 30, 'thirty-one': 31, 'thirty-two': 32, 'thirty-three': 33,
    'thirty-four': 34, 'thirty-five': 35, 'thirty-six': 36, 'thirty-seven': 37,
    'thirty-eight': 38, 'thirty-nine': 39, 'forty': 40, 'forty-one': 41,
    'forty-two': 42, 'forty-three': 43, 'forty-four': 44, 'forty-five': 45,
    'forty-six': 46, 'forty-seven': 47, 'forty-eight': 48, 'forty-nine': 49,
    'fifty': 50, 'fifty-one': 51, 'fifty-two': 52, 'fifty-three': 53,
    'fifty-four': 54, 'fifty-five': 55, 'fifty-six': 56, 'fifty-seven': 57,
    'fifty-eight': 58, 'fifty-nine': 59, 'sixty': 60, 'sixty-one': 61,
    'sixty-two': 62, 'sixty-three': 63, 'sixty-four': 64, 'sixty-five': 65,
    'sixty-six': 66, 'sixty-seven': 67, 'sixty-eight': 68, 'sixty-nine': 69,
    'seventy': 70, 'seventy-one': 71, 'seventy-two': 72, 'seventy-three': 73,
    'seventy-four': 74, 'seventy-five': 75, 'seventy-six': 76, 'seventy-seven': 77,
    'seventy-eight': 78, 'seventy-nine': 79, 'eighty': 80, 'eighty-one': 81,
    'eighty-two': 82, 'eighty-three': 83, 'eighty-four': 84, 'eighty-five': 85,
    'eighty-six': 86, 'eighty-seven': 87, 'eighty-eight': 88, 'eighty-nine': 89,
    'ninety': 90, 'ninety-one': 91, 'ninety-two': 92, 'ninety-three': 93,
    'ninety-four': 94, 'ninety-five': 95, 'ninety-six': 96, 'ninety-seven': 97,
    'ninety-eight': 98, 'ninety-nine': 99, 'hundred': 100
}



# Function to process the string with regex and apply transformations
def process_string(text):
    
    # Regular expression to match "word and word" pattern
    pattern = re.compile(r'(\b\w+\b) and (\b\w+\b)')
    
    # Reverse dictionary to convert integer values back to words
    num_to_word = {v: k for k, v in word_to_num.items()}

    # Function to convert word number to integer
    def word_to_int(word):
        return word_to_num.get(word.lower(), None)

    # Function to convert integer to word
    def int_to_word(num):
        return num_to_word.get(num, None)
    
    def replace_match(match):
        word1 = match.group(1)
        word2 = match.group(2)
        # Convert words to their numeric values
        num1 = word_to_int(word1)
        num2 = word_to_int(word2)
        
        # If both numbers are less than 9, return them as word form
        if (num1 is not None and num1 < 9) and (num2 is not None and num2 < 9):
            return f"{word1} and {word2}"  # No change if both are < 9
        
        # If either number is greater than or equal to 9, convert to numeric form
        if (num1 is not None and num1 >= 9) or (num2 is not None and num2 >= 9):
            # Convert both to numeric form
            num1 = num1 if num1 is not None else word1
            num2 = num2 if num2 is not None else word2
            return f"{num1} and {num2}"  # Replace with numeric values
        
        return match.group(0)  # Return the match as is if both are < 9
    
    # Apply regex substitution with the replace function
    return pattern.sub(replace_match, text)


def write_to_log(doc_id, user):
    global global_logs
    # output_dir = os.path.join('output', str(doc_id))
    current_date = datetime.now().strftime("%Y-%m-%d")
    output_path_file = Path(os.getcwd()) / 'output' / user / current_date / str(doc_id) / 'text'
    os.makedirs(output_path_file, exist_ok=True)
    
    log_file_path = os.path.join(output_path_file, 'global_logs.txt')
    with open(log_file_path, 'a', encoding='utf-8') as log_file:
        log_file.write("\n".join(global_logs))
    global_logs = []
    



def move_punctuation_inside_quotes(runs, line_number_offset):
    # reposition_pattern = re.compile(r'([‘“"])(.*?)(["’”])([?!])')
    reposition_pattern = re.compile(r'([‘“"\'])(.*?)(["’”\'])([?!])')
    def reposition_replacer(match):
        opening_quote = match.group(1)
        content = match.group(2)
        closing_quote = match.group(3)
        punctuation = match.group(4)
        return f"{opening_quote}{content}{punctuation}{closing_quote}"
    
    for run in runs:
        run.text = reposition_pattern.sub(reposition_replacer, run.text)
        run.text = re.sub(r'(?<=[?!][’”"\'])\.', '', run.text)


def single_to_double_quotes(runs):
    single_quote_index = -1  # Store the index of the first single quote
    single_quote_run = -1  # Store the run index of the first single quote

    for i, run in enumerate(runs):
        text = run.text

        # Find exact index of `'` and `.'` in the text
        if single_quote_index == -1:  # Store first occurrence only
            single_quote_index = text.find("‘")
            if single_quote_index != -1:
                single_quote_run = i  # Store the run where it occurred

        dot_single_quote_index = text.find(".’")

        if dot_single_quote_index != -1 and single_quote_index != -1:
            # Replace single quote `'` with `"` at single_quote_index
            runs[single_quote_run].text = (
                runs[single_quote_run].text[:single_quote_index] + '“' +
                runs[single_quote_run].text[single_quote_index + 1:]
            )

            # Replace `.'` with `."` in the current run
            runs[i].text = (
                text[:dot_single_quote_index] + '.”' + text[dot_single_quote_index + 2:]
            )

            # Reset tracking variables
            single_quote_index = -1
            single_quote_run = -1



def elide_consecutive_references(runs):
    """
    Processes a list of Word document runs and converts consecutive figure, table, or equation references
    into an elided range format.
    
    For example, this converts:
       "Figures 1.1, 1.2, 1.3, 1.4 and 1.5"
    into:
       "Figures 1.1–1.5"
    """
    if not runs:  # Check if runs list is empty
        return  

    # Matches references like "Figures 1.1, 1.2, 1.3, 1.4 and 1.5"
    pattern = r"\b(Figures|Tables|Equations|Parts|Figure) ((?:\d+\.\d+(?:, | and )?)+)"

    # Collect text from runs
    full_text = "".join(run.text for run in runs)

    def replace_match(match):
        prefix = match.group(1)  # e.g., 'Figures'
        numbers = re.split(r", | and ", match.group(2))  # e.g., ["1.1", "1.2", "1.3", "1.4", "1.5"]
        processed_numbers = []

        # Convert numbers into tuples (major, minor), e.g. (1, 1)
        parsed_numbers = [tuple(map(int, num.split("."))) for num in numbers]

        # Identify consecutive ranges
        start = parsed_numbers[0]
        end = start

        for i in range(1, len(parsed_numbers)):
            # Check if the current number is consecutive to the previous
            if parsed_numbers[i][0] == end[0] and parsed_numbers[i][1] == end[1] + 1:
                end = parsed_numbers[i]  # Extend range
            else:
                # Append previous range
                if start == end:
                    processed_numbers.append(f"{start[0]}.{start[1]}")
                else:
                    processed_numbers.append(f"{start[0]}.{start[1]}–{end[0]}.{end[1]}")
                start = end = parsed_numbers[i]

        # Append the last identified range
        if start == end:
            processed_numbers.append(f"{start[0]}.{start[1]}")
        else:
            processed_numbers.append(f"{start[0]}.{start[1]}–{end[0]}.{end[1]}")

        return f"{prefix} {', '.join(processed_numbers)}"

    # Replace text in the full_text
    new_text = re.sub(pattern, replace_match, full_text)

    # Update the runs: place the new_text in the first run and clear the others
    runs[0].text = new_text
    for i in range(1, len(runs)):
        runs[i].text = ""


def units_with_bracket(runs, used_units):
    units = {
        "s": "second",
        "m": "meter",
        "kg": "kilogram",
        "A": "ampere",
        "K": "kelvin",
        "mol": "mole",
        "cd": "candela"
    }
    
    
    pattern = r'\b(\d+(?:\.\d+)?)\s+(' + '|'.join(re.escape(unit) for unit in units.keys()) + r')\b'
    
    for run in runs:
        def replace_unit(match):
            number = match.group(1)
            unit = match.group(2)
            if unit in used_units:
                return match.group(0)  # If already replaced, return the original "number unit"
            else:
                used_units.add(unit)
                full_form = units[unit]
                # Append an "s" to the full form if needed (except for "mol" or if already plural)
                if unit != "mol" and not full_form.endswith("s"):
                    full_form += "s"
                replacement = f"{number} {full_form} ({unit.lower()})"
                global_logs.append(f"Unit {unit} replaced with {replacement}")
                return replacement

        run.text = re.sub(pattern, replace_unit, run.text)



def process_doc_function1(payload: dict, doc: Document, doc_id, user):
    """
    This function processes the document by converting century notations
    and highlighting specific words.
    """
    line_number = 1
    abbreviation_dict = fetch_abbreviation_mappings()
    used_units = set()
    seen_symbols = set()
    replaced_units = set()

    for para in doc.paragraphs:
        # if para.style.name.startswith("Heading") and para.text.strip().startswith("Chapter"):
        #     used_units = set()
        # insert_thin_space_between_number_and_unit_in_runs(para.runs, line_number)
        # elide_consecutive_references(para.runs)
        
        single_to_double_quotes(para.runs)
        move_punctuation_inside_quotes(para.runs, line_number)
        apply_abbreviation_mapping(para.runs, abbreviation_dict, line_number)
        convert_century_in_runs(para.runs, line_number)
        set_latinisms_to_roman_in_runs(para.runs, line_number)
        process_symbols_mark_in_runs(para.runs, line_number, seen_symbols)
        apply_remove_italics_see_rule_in_runs(para.runs)
        set_number_to_no_in_runs(para.runs, line_number)
        format_titles_us_english_with_logging_in_runs(para.runs, line_number)
        enforce_am_pm_in_runs(para.runs, line_number)
        enforce_serial_comma_in_runs(para.runs)
        if(payload["2"] == False):
            rename_section_in_runs(para.runs)
        replace_ampersand_in_runs(para.runs, line_number)
        correct_possessive_names_in_runs(para.runs, line_number)
        units_with_bracket(para.runs, used_units)
        remove_and_in_runs(para.runs, line_number)
        remove_quotation_in_runs(para.runs, line_number)
        correct_acronyms_in_runs(para.runs, line_number)
        enforce_eg_rule_with_logging_in_runs(para.runs, line_number)
        enforce_ie_rule_with_logging_in_runs(para.runs, line_number)
        standardize_etc_in_runs(para.runs, line_number)
        process_paragraph(para.runs, line_number)
        line_number += 1

    write_to_log(doc_id, user)

