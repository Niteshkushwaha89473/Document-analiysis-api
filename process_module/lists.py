import os
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn

global_logs = []



# def is_list_item(para):
#     """Check if paragraph is part of a list structure using XPath."""
#     numPr = para._p.xpath('.//w:numPr')
#     return len(numPr) > 0


# def get_list_level(para):
#     """Get the nesting level of list items (0 = top level)."""
#     numPr = para._p.xpath('.//w:numPr')
#     if not numPr:
#         return None
#     ilvl = numPr[0].xpath('.//w:ilvl/@w:val')
#     return int(ilvl[0]) if ilvl else 0


# def convert_nested_lists(doc):
#     """
#     Convert nested lists (level 1+) to dash lists while preserving:
#     - Parent list formatting
#     - Text styles
#     - Original indentation
#     """
#     for para in doc.paragraphs:
#         if not is_list_item(para):
#             continue
            
#         level = get_list_level(para)
#         if level is None or level < 1:
#             continue
            
#         for numPr in para._p.xpath('.//w:numPr'):
#             parent = numPr.getparent()
#             if parent is not None:
#                 parent.remove(numPr)
                
#         indent = ' ' * 4 * (level - 1)  # Adjusted for 0-based level
#         dash = f'{indent}- '
        
#         # Preserve existing formatting while adding dash
#         if para.runs:
#             first_run = para.runs[0]
#             original_text = first_run.text or ''
#             cleaned_text = original_text.lstrip('. \t')
#             first_run.text = dash + cleaned_text
#         else:
#             para.add_run(dash)


from docx import Document
from docx.oxml.ns import qn

def is_list_item(para):
    """Check if paragraph is part of a list structure."""
    pPr = para._p.pPr
    return pPr is not None and pPr.numPr is not None

def get_list_level(para):
    """Get the nesting level of list items (0 = top level)."""
    numPr = para._p.xpath('.//w:numPr')
    if not numPr:
        return None
    ilvl = numPr[0].xpath('.//w:ilvl/@w:val')
    return int(ilvl[0]) if ilvl else 0

def convert_nested_lists(doc):
    """
    Convert nested lists (level 1+) to dash lists while preserving:
    - Parent list formatting
    - Text styles
    - Original indentation

    For nested list items (level 1+), if the most recent parent list item (level 0)
    contains the word "steps" (case-insensitive), then the sublist items are numbered
    sequentially (e.g. "1. ", "2. ", etc.) instead of using a dash.
    """
    
    current_steps_flag = False
    numbering_counter = 1
    
    for para in doc.paragraphs:
        if not is_list_item(para):
            continue
        
        level = get_list_level(para)
        if level is None:
            continue

        if level == 0:
            if "steps" in para.text.lower():
                current_steps_flag = True
            else:
                current_steps_flag = False
            numbering_counter = 1
            continue
        
        pPr = para._p.get_or_add_pPr()
        if pPr is not None and pPr.numPr is not None:
            pPr.remove(pPr.numPr)
        
        indent = ' ' * 4 * (level - 1)
        
        if current_steps_flag:
            prefix = f'{indent}{numbering_counter}. '
            numbering_counter += 1
        else:
            prefix = f'{indent}- '
        
        if para.runs:
            first_run = para.runs[0]
            if first_run.text:
                cleaned_text = first_run.text.lstrip('. \t')
                first_run.text = prefix + cleaned_text
            else:
                first_run.text = prefix
        else:
            para.add_run(prefix)
    
    # doc.save(output_path)

# Usage example:
# convert_nested_lists("testing1.docx", "output2.docx")

import re
from docx import Document  # assuming you're using python-docx

# def int_to_roman(num):
#     """Convert an integer to a Roman numeral."""
#     val = [
#         1000, 900, 500, 400,
#         100, 90, 50, 40,
#         10, 9, 5, 4,
#         1
#     ]
#     syms = [
#         "M", "CM", "D", "CD",
#         "C", "XC", "L", "XL",
#         "X", "IX", "V", "IV",
#         "I"
#     ]
#     roman_num = ''
#     for i, v in enumerate(val):
#         while num >= v:
#             roman_num += syms[i]
#             num -= v
#     return roman_num.lower()

# def roman_to_int(s):
#     """Convert a Roman numeral to an integer."""
#     roman_map = {'I': 1, 'V': 5, 'X': 10, 'L': 50, 'C': 100, 'D': 500, 'M': 1000}
#     num = 0
#     s = s.upper()
#     for i in range(len(s)):
#         # if the current value is less than the next one, subtract it
#         if i + 1 < len(s) and roman_map[s[i]] < roman_map[s[i+1]]:
#             num -= roman_map[s[i]]
#         else:
#             num += roman_map[s[i]]
#     return num

# def swap_number_roman(match):
#     """Replace (number) with (roman) and vice versa."""
#     token = match.group(1)
#     if token.isdigit():
#         # Convert number to Roman numeral.
#         return f"({int_to_roman(int(token))})"
#     else:
#         # Convert Roman numeral to number.
#         try:
#             num = roman_to_int(token)
#             return f"({num})"
#         except Exception:
#             return match.group(0)

# def process_paragraph(paragraph_text):
#     """
#     If the paragraph contains 'steps:', swap all (number) with (roman numeral) 
#     and all (roman numeral) with (number).
#     """
#     if "steps:" in paragraph_text:
#         # This pattern matches tokens like (1) or (i)
#         pattern = re.compile(r"\((\d+|[ivxlcdmIVXLCDM]+)\)")
#         return pattern.sub(swap_number_roman, paragraph_text)
#     return paragraph_text

# def roman_to_int_list(doc):
#     """
#     Process each paragraph in the document.
#     If a paragraph contains the word 'steps:', then swap the numbering styles.
#     """
#     for para in doc.paragraphs:
#         new_text = process_paragraph(para.text)
#         # Update paragraph text only if changes occurred.
#         if new_text != para.text:
#             para.text = new_text



def int_to_roman(num):
    """Convert an integer to a Roman numeral."""
    val = [
        1000, 900, 500, 400,
        100, 90, 50, 40,
        10, 9, 5, 4,
        1
    ]
    syms = [
        "M", "CM", "D", "CD",
        "C", "XC", "L", "XL",
        "X", "IX", "V", "IV",
        "I"
    ]
    roman_num = ''
    for i, v in enumerate(val):
        while num >= v:
            roman_num += syms[i]
            num -= v
    return roman_num.lower()

def roman_to_int(s):
    """Convert a Roman numeral to an integer."""
    roman_map = {'I': 1, 'V': 5, 'X': 10, 'L': 50, 'C': 100, 'D': 500, 'M': 1000}
    num = 0
    s = s.upper()
    for i in range(len(s)):
        # if the current value is less than the next one, subtract it
        if i + 1 < len(s) and roman_map[s[i]] < roman_map[s[i+1]]:
            num -= roman_map[s[i]]
        else:
            num += roman_map[s[i]]
    return num

def swap_number_roman(match):
    """Replace (number) with (roman) and vice versa."""
    token = match.group(1)
    if token.isdigit():
        # Convert number to Roman numeral.
        return f"({int_to_roman(int(token))})"
    else:
        # Convert Roman numeral to number.
        try:
            num = roman_to_int(token)
            return f"({num})"
        except Exception:
            return match.group(0)

def process_paragraph(paragraph):
    """
    If the paragraph contains 'steps:', process each run to swap (number) with (roman numeral)
    and vice versa, preserving the original formatting.
    """
    if "steps:" in paragraph.text:
        # Define pattern to match tokens like (1) or (i)
        pattern = re.compile(r"\((\d+|[ivxlcdmIVXLCDM]+)\)")
        for run in paragraph.runs:
            # Only modify runs that contain potential tokens.
            if pattern.search(run.text):
                run.text = pattern.sub(swap_number_roman, run.text)

def roman_to_int_list(doc):
    """
    Process each paragraph in the document.
    If a paragraph contains the word 'steps:', then swap the numbering styles in each run.
    """
    for para in doc.paragraphs:
        process_paragraph(para)





def write_to_log(doc_id, user):
    global global_logs
    current_date = datetime.now().strftime("%Y-%m-%d")
    output_path_file = Path(os.getcwd()) / 'output' / user / current_date / str(doc_id) / 'text' 
    # dir_path = output_path_file.parent

    # output_dir = os.path.join('output', str(doc_id))
    os.makedirs(output_path_file, exist_ok=True)
    log_file_path = os.path.join(output_path_file, 'global_logs.txt')

    with open(log_file_path, 'a', encoding='utf-8') as log_file:
        log_file.write("\n".join(global_logs))

    global_logs = []




def process_doc_function11(payload: dict, doc: Document, doc_id, user):
    convert_nested_lists(doc)
    roman_to_int_list(doc)
    # for para in doc.paragraphs:
    #     print('hello')
        
    write_to_log(doc_id, user)